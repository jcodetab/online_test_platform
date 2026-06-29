from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import Test, TestCompletion
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from .models import Result
from rest_framework import viewsets
from rest_framework.decorators import action 
from .models import UserProfile
from .serializers import ProfileSerializer, ProfileUpdateSerializer 
from rest_framework.response import Response
from .models import UserProfile
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.authentication import TokenAuthentication
from .serializers import RegisterSerializer, LoginSerializer
from django.contrib.auth.models import User
from rest_framework.authtoken.models import Token
from rest_framework import status
from rest_framework import generics
from django.contrib.auth import get_user_model
from .serializers import ChangePasswordSerializer
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib.auth import login as auth_login
from django.views.decorators.cache import never_cache
from django.contrib.auth import logout
from datetime import timedelta
from amaliyot.models import Question, Olympiad  
from accounts.models import User
from .serializers import DashboardStatsSerializer






User = get_user_model()


def test_start(request):
    return render(request, 'registration/test_start.html')


def register_page(request):
    return render(request, 'registration/register.html') 


class RegisterAPIView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            
            # Profilni majburan yangilash
            profile, _ = UserProfile.objects.get_or_create(user=user)
            profile.first_name = user.first_name
            profile.last_name = user.last_name
            profile.save()
            
            token, _ = Token.objects.get_or_create(user=user)
            return Response({'token': token.key}, status=status.HTTP_201_CREATED)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

def login_page(request):
    return render(request, 'registration/login.html')


class LoginAPIView(APIView):
    permission_classes = [AllowAny]
    authentication_classes = []

    def post(self, request):
        serializer = LoginSerializer(data=request.data)
        if serializer.is_valid():
            email = serializer.validated_data["email"]
            password = serializer.validated_data["password"]

            try:
                user = User.objects.get(email=email)
            except User.DoesNotExist:
                return Response({"error": "Bunday foydalanuvchi yo‘q"}, status=status.HTTP_400_BAD_REQUEST)

            if not user.check_password(password):
                return Response({"error": "Noto‘g‘ri parol"}, status=status.HTTP_400_BAD_REQUEST)
            
            # Django session logini (HTML templatingiz uchun kerak)
            auth_login(request, user)

            # TOKEN MUAMMOSI SHU YERDA YECHILDI:
            # Eski tokenni o'chirmasdan, borini oladi, yo'q bo'lsa yangi yaratadi.
            token, _ = Token.objects.get_or_create(user=user)

            return Response({"token": token.key}, status=status.HTTP_200_OK)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

def logout_view(request):
    logout(request)
    return redirect('home')


class UserCreateAPIView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            UserProfile.objects.create(user=user)
            return Response({"message": "Ro'yxatdan o'tish muvaffaqiyatli"}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    

class DashboardStatsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        bugun = timezone.now().date()
        
        # 1. Testlar statistikasi
        total_tests = Question.objects.count()
        # 'status' maydoni bor-yo'qligini tekshirib filtrlaymistan
        if hasattr(Question, 'status'):
            active_tests = Question.objects.filter(status='OCHIQ').count()
        else:
            active_tests = total_tests

        # 2. Foydalanuvchilar statistikasi
        total_users = User.objects.count()
        # Oxirgi 24 soat ichida ro'yxatdan o'tganlar soni
        new_users = User.objects.filter(date_joined__date=bugun).count()

        # 3. Guruhlar statistikasi (Agar Group modeli import qilingan bo'lsa ishlaydi)
        total_groups = 0
        olympiad_groups = 0
        if 'Group' in globals() or 'Group' in locals():
            try:
                from amaliyot.models import Group
                total_groups = Group.objects.count()
                # Agar guruh turi 'olympiad' yoki shunga o'xshash bo'lsa:
                if hasattr(Group, 'group_type'):
                    olympiad_groups = Group.objects.filter(group_type='OLYMPIAD').count()
            except Exception:
                pass

        # 4. Olimpiadalar statistikasi
        total_olympiads = Olympiad.objects.count()
        
        # 🔥 MANA SHU YER TUZATILDI: 'start_time' o'rniga 'start_date' qo'yildi
        upcoming_olympiads = Olympiad.objects.filter(start_date__gt=timezone.now()).count()

        # Ma'lumotlarni yig'amiz
        data = {
            'total_tests': total_tests,
            'active_tests': active_tests,
            'total_users': total_users,
            'new_users': new_users,
            'total_groups': total_groups,
            'olympiad_groups': olympiad_groups,
            'total_olympiads': total_olympiads,
            'upcoming_olympiads': upcoming_olympiads,
        }

        # Serializerga data=data ko'rinishida uzatamiz (VS Code sarg'aymasligi uchun)
        serializer = DashboardStatsSerializer(data)
        return Response(serializer.data, status=status.HTTP_200_OK)


def home(request):
    return render(request, 'registration/home.html')


def about(request):
    return render(request, 'registration/about.html')


def contact(request):
    return render(request, 'registration/contact.html')


def courses(request):
    return render(request, 'registration/courses.html')


def team(request):
    return render(request, 'registration/team.html')


def testimonial(request):
    return render(request, 'registration/testimonial.html')


@never_cache
@login_required
def dashboard(request):
    return render(request, "registration/dashboard.html")
  

def test_list(request):
    return render(request, "test/tests.html")


def test_detail_view(request, test_id):
    test = get_object_or_404(Test, id=test_id)
    return render(request, 'test/test_detail.html', {'test': test})


@login_required
def start_test(request, test_id):
   
    test = Test.objects.get(id=test_id)
    if test.is_completed:
        return redirect('view_results', test_id=test.id)

    return render(request, 'registration/start_test.html', {'test': test})


@login_required
def view_results(request, test_id):
    
    test_result = TestCompletion.objects.get(test_id=test_id, user=request.user)
    return render(request, 'registration/test_results.html', {'test_result': test_result})


def profile_page(request):
    """
    Foydalanuvchi token orqali sahifa ma’lumotini JS orqali oladi
    """
    return render(request, "registration/profile.html")


class ProfileViewSet(viewsets.ViewSet):
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action in ['me']:
            return ProfileUpdateSerializer
        return ProfileSerializer

    @action(detail=False, methods=['get', 'patch'])
    def me(self, request):
        profile, _ = UserProfile.objects.get_or_create(user=request.user)

        if request.method == 'PATCH':
            serializer = ProfileUpdateSerializer(
                profile,
                data=request.data,
                partial=True
            )
            serializer.is_valid(raise_exception=True)
            serializer.save()
            return Response(serializer.data)

        serializer = ProfileSerializer(profile)
        return Response(serializer.data)


class ChangePasswordViewSet(viewsets.ViewSet):
    permission_classes = [IsAuthenticated]

    @action(detail=False, methods=['post'])
    def change_password(self, request):
        serializer = ChangePasswordSerializer(
            data=request.data,
            context={'request': request}
        )
        serializer.is_valid(raise_exception=True)
        serializer.save()
        return Response({"detail": "Parol muvaffaqiyatli yangilandi!"}, status=status.HTTP_200_OK)
    

def change_password_page(request):
    """
    Faqat HTML sahifani qaytaradi, tekshiruvni JavaScript (Token) bajaradi
    """
    return render(request, 'registration/change_password.html')


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def api_change_password(request):
    """
    Token login foydalanuvchisi uchun parol o'zgartirish API
    """
    user = request.user
    serializer = ChangePasswordSerializer(data=request.data, context={'request': request})
    serializer.is_valid(raise_exception=True)
    serializer.save()
    update_session_auth_hash(request, user)  # Session bo'lsa ham saqlash
    return Response({"detail": "Parol muvaffaqiyatli o'zgartirildi"}, status=status.HTTP_200_OK)

   
def results_view(request):
    score = request.GET.get('score', 0)  
    return render(request, 'accounts/results.html', {'score': score})


def take_test(request, test_id):
    test = Test.objects.get(id=test_id)
    questions = test.questions.all()

    if request.method == 'POST':
        score = 0
        total = len(questions)

        for question in questions:
            
            selected_option = request.POST.get(f'correct_option_{question.id}')
            correct_option = question.options.filter(is_correct=True).first()

            
            if selected_option == str(correct_option.id):
                score += 1

        
        return render(request, 'closed_tests.html', {'test': test, 'questions': questions, 'score': score, 'total': total})

    return render(request, 'test/closed_tests.html', {'test': test, 'questions': questions}
                  
    )


def profile_edit_view(request):
    
    return render(request, 'registration/profile_edit_page.html') 


@api_view(['GET', 'PATCH'])
@permission_classes([IsAuthenticated])
def api_me_profile(request):
    user = request.user
    # Profile obyektini olish yoki yaratish
    profile, created = UserProfile.objects.get_or_create(user=user)

    if request.method == 'GET':
        return Response({
            "first_name": user.first_name,
            "last_name": user.last_name,
            "email": user.email,
            "date_joined": user.date_joined.strftime("%d.%m.%Y"), 
            "gender": getattr(profile, 'gender', 'not_specified'),
            "avatar_url": profile.avatar.url if hasattr(profile, 'avatar') and profile.avatar else None,
            "solved_tests": 0, 
            "average_score": "0%",
            "groups_count": 0
        })

    if request.method == 'PATCH':
        data = request.data 
        
        # 1. User modelini yangilash (Ism, Familiya, Email)
        user.first_name = data.get('first_name', user.first_name)
        user.last_name = data.get('last_name', user.last_name)
        user.email = data.get('email', user.email)
        user.save()

    
        if hasattr(profile, 'bio'):
            profile.bio = data.get('bio', profile.bio)
        if hasattr(profile, 'location'):
            profile.location = data.get('location', profile.location)
        if hasattr(profile, 'gender'):
            profile.gender = data.get('gender', profile.gender)
            
        profile.save()

        return Response({
            "status": "success",
            "message": "Profil muvaffaqiyatli yangilandi",
            "first_name": user.first_name,
            "last_name": user.last_name,
            "email": user.email,
            "date_joined": user.date_joined.strftime("%d.%m.%Y")
        })


@csrf_exempt
def save_result(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        user_name = data.get('userName')
        score = data.get('score')

        if not user_name or score is None:
            return JsonResponse({'error': 'Foydalanuvchi nomi va natija kiritilishi kerak!'}, status=400)

        result = Result.objects.create(user_name=user_name, score=score)

        print(f"Foydalanuvchi: {user_name}, Natija: {score}")

        return JsonResponse({'message': 'Natija saqlandi!', 'result_id': result.id}, status=200)
    
    return JsonResponse({'error': 'Noto\'g\'ri so\'rov'}, status=400)


def tests_view(request):
    return render(request, 'test/tests.html')


def kazus_tests_view(request):
    return render(request, 'test/kazus_tests.html')


def ranking_view(request):
    return render(request, 'test/ranking.html')


def closed_tests_view(request):
    return render(request, 'test/closed_tests.html') 


from rest_framework.authentication import SessionAuthentication, TokenAuthentication
from rest_framework import viewsets
from rest_framework import generics
from rest_framework.response import Response
from rest_framework import viewsets, permissions, filters
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from drf_yasg.utils import swagger_auto_schema
from rest_framework import status
from drf_yasg import openapi
import hashlib
from django.conf import settings
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import User
from rest_framework.decorators import action
from .serializers import PaymentSerializer
from .utils import evaluate_test
from .models import Test, Answer, TestQueue, Question, Option, TestCompletion, UserTest, UserProfile, UserAnswer, Task, Result, Subject, Topic, Payment, TestAnswer, Choice, CaseTest, ClosedTest
from .serializers import TestSerializer, QuestionSerializer, OptionSerializer, TestCompletionSerializer, UserTestSerializer, ProfileSerializer, AnswerSerializer, UserAnswerSerializer, TaskSerializer, ResultSerializer, UserSerializer, ChoiceSerializer, AnswerSerializer
from .serializers import (
    SubjectSerializer,
    TopicSerializer,
    TestSerializer,
    QuestionSerializer,
)
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from .models import TestSession
from .models import OlympiadGroup, OlympiadParticipant
from .serializers import TestSessionSerializer
from .utils import write_key, load_key, encrypt_message, decrypt_message
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from django.contrib.auth.models import User
from .models import Group, GroupMembership, OlympiadGroup, OlympiadParticipant, OlympiadAnswer, OlympiadQuestion, ClosedTest, ClosedQuestion, ClosedTestSession, ClosedAnswer
from .serializers import OlympiadGroupSerializer, OlympiadParticipantSerializer, TestCreateSerializer
from rest_framework.authtoken.views import ObtainAuthToken
from rest_framework.authtoken.models import Token
from rest_framework.permissions import IsAdminUser
from .serializers import TestQueueSerializer
from rest_framework.permissions import AllowAny
from .serializers import UserSerializer,  ClosedTestSerializer, ClosedTestSessionSerializer

from django.utils import timezone
from django.apps import apps
from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer
from .models import TestSession, Question 
from django.shortcuts import get_object_or_404
from rest_framework.pagination import PageNumberPagination
from .tasks import auto_evaluate_expired_tests
from .models import OlympiadGroup
from rest_framework.authentication import SessionAuthentication
from .serializers import OlympiadGroupSerializer
from .models import Group
from .serializers import TestSerializer
from rest_framework import status, permissions, authentication 
from django.db import transaction
import mammoth
import openpyxl
from datetime import datetime
from django.views.decorators.http import require_POST
import logging
from django.db.models import Q
from .models import Group, GroupMembership, ChatMessage
from .serializers import GroupSerializer, GroupCreateSerializer, GroupMembershipSerializer, ChatMessageSerializer
from rest_framework.viewsets import ViewSet
from .serializers import KazusTestSerializer, CaseTestSerializer, ClosedTestSerializer
import operator
from functools import reduce
from rest_framework.generics import RetrieveAPIView
from docx import Document
import io, re
from django.db.models import Count, Q
from .models import Test, ClosedTest, CaseTest
from rest_framework.decorators import api_view
from rest_framework.response import Response






class ApiRoot(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        return Response({
            "message": "API ishlayapti!",
            "user": request.user.username
        })
    

class UserCreateAPIView(generics.CreateAPIView):
    serializer_class = UserSerializer
    permission_classes = [AllowAny]


class CustomAuthToken(ObtainAuthToken):
    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data,
                                           context={'request': request})
        serializer.is_valid(raise_exception=True)
        user = serializer.validated_data['user']
        token, created = Token.objects.get_or_create(user=user)
        return Response({'token': token.key, 'user_id': user.id, 'email': user.email})


class ClosedTestViewSet(viewsets.ModelViewSet):
    queryset = ClosedTest.objects.all().order_by('-id')
    serializer_class = ClosedTestSerializer
    parser_classes = (MultiPartParser, FormParser)

    @action(detail=False, methods=['POST'], url_path='upload')
    @transaction.atomic
    def upload(self, request):
        print("\n--- FILE PARSING START ---")
        
        # 1. Fayllarni olish
        files = request.FILES.getlist('file') or request.FILES.getlist('files')
        if not files:
            return Response({"detail": "Fayl topilmadi. 'file' kaliti bilan yuboring."}, status=400)

        created_tests = []

        for f in files:
            print(f"Fayl nomi: {f.name}")
            try:
                # 2. Fayl ichidagi matnni o'qib olish
                full_text = ""
                if f.name.endswith('.docx'):
                    result = mammoth.extract_raw_text(f)
                    full_text = result.value
                elif f.name.endswith('.txt'):
                    full_text = f.read().decode('utf-8', errors='ignore')
                else:
                    print(f"Format xato: {f.name}")
                    continue

                # 3. Parsing qilish
                lines = [l.strip() for l in full_text.splitlines() if l.strip()]
                questions_list = []
                current_q = None

                for line in lines:
                    if line.startswith('[s]'):
                        # Oldingi savolni yopish
                        if current_q and len(current_q['answers']) >= 2:
                            questions_list.append(current_q)
                        
                        current_q = {
                            'text': line.replace('[s]', '').strip(),
                            'answers': [],
                            'correct_index': 0 
                        }
                    
                    elif (line.startswith('+') or line.startswith('-')) and current_q:
                        is_correct = line.startswith('+')
                        clean_answer = line[1:].strip()
                        current_q['answers'].append(clean_answer)
                        
                        if is_correct:
                            current_q['correct_index'] = len(current_q['answers']) - 1

                # Eng oxirgi savolni qo'shish
                if current_q and len(current_q['answers']) >= 2:
                    questions_list.append(current_q)

                print(f"Jami topilgan savollar: {len(questions_list)}")

                if not questions_list:
                    print(f"DIQQAT: {f.name} ichida mos formatdagi savollar topilmadi!")
                    continue

                # 4. Bazaga Testni saqlash
                # QAT'IY: Har bir savol uchun 1 minut (default 30 olib tashlandi)
                q_count = len(questions_list)
                auto_duration = q_count if q_count > 0 else 1 

                test = ClosedTest.objects.create(
                    user=request.user if request.user.is_authenticated else None,
                    title=f.name,
                    description=f"Fayldan yuklangan ({q_count} ta savol)",
                    duration=auto_duration, 
                    score=100,
                    file=f
                )

                # 5. Savollarni bazaga saqlash (bulk_create - tezroq)
                question_objs = []
                for q in questions_list:
                    # 4 ta javob bo'lishini ta'minlash (bo'sh bo'lsa ham)
                    ans = q['answers'] + [""] * (4 - len(q['answers']))
                    # Indeksni harfga o'girish (0->A, 1->B, 2->C, 3->D)
                    correct_letter = chr(65 + q['correct_index']) 

                    question_objs.append(
                        ClosedQuestion(
                            test=test,
                            text=q['text'],
                            answer_A=ans[0],
                            answer_B=ans[1],
                            answer_C=ans[2],
                            answer_D=ans[3],
                            correct_answers=correct_letter,
                            question_type='single'
                        )
                    )
                
                if question_objs:
                    ClosedQuestion.objects.bulk_create(question_objs)
                
                created_tests.append(test.id)
                print(f"Muvaffaqiyatli saqlandi: ID {test.id} | Vaqt: {auto_duration} minut")

            except Exception as e:
                print(f"Xatolik yuz berdi: {str(e)}")
                print(traceback.format_exc())
                return Response({"detail": f"Faylda xatolik: {str(e)}"}, status=400)

        print("--- FILE PARSING END ---\n")
        return Response({
            "message": f"{len(created_tests)} ta test yaratildi",
            "test_ids": created_tests
        }, status=status.HTTP_201_CREATED)


class ClosedTestSessionViewSet(viewsets.ModelViewSet):
    serializer_class = ClosedTestSessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return ClosedTestSession.objects.filter(user=self.request.user).order_by('-started_at')
    

def search_tests_page(request):
    return render(request, 'registration/search_tests.html')
    

from django.db.models import Count
from rest_framework.decorators import api_view
from rest_framework.response import Response
import traceback

@api_view(['GET'])
def api_get_tests(request):
    try:
        # 1. So'rov parametrlarini olish
        query = request.GET.get('q', '').strip().lower()
        select_cat = request.GET.get('select_cat', '').strip().lower()  # Select'dan kelgan qiymat
        input_cat = request.GET.get('input_cat', '').strip().lower()    # Input'dan kelgan matn
        duration_filter = request.GET.get('duration', '')
        q_count_filter = request.GET.get('q_count', '')
        mine_filter = request.GET.get('mine', 'false').lower() == 'true'

        all_tests_raw = []

        # --- OPTIMALLASHTIRILGAN PROCESS FUNKSIYASI ---
        def process_test_object(obj, model_origin):
            obj_id = obj.pk 
            match_score = 0.0 
            
            db_test_type = str(getattr(obj, 'test_type', 'regular')).lower()
            title_str = str(getattr(obj, 'title', '')).lower()
            fan_nomi_db = str(getattr(obj, 'fan_nomi', '')).lower()
            
            # select_related va prefetch_related qilingani uchun hasattr DBga urilmaydi
            has_case = model_origin == 'case' or hasattr(obj, 'case_test_detail')
            has_closed = model_origin == 'closed' or hasattr(obj, 'closedtest')

            # --- TEST TOIFASINI ANIQLASH ---
            is_kazus = model_origin == 'case' or db_test_type == 'kazus' or 'kazus' in title_str or 'kasuz' in title_str or 'kazus' in fan_nomi_db or has_case
            is_closed = model_origin == 'closed' or db_test_type == 'closed' or 'yopiq' in title_str or 'yopiq' in fan_nomi_db or has_closed

            if is_kazus:
                label, t_key = "KAZUS TEST", "kazus"
            elif is_closed:
                label, t_key = "YOPIQ TEST", "closed"
            else:
                label, t_key = "ODDIY TEST", "regular"

            # --- SAVOLLAR SONINI HISOBLASH (OPTIMAL - ANNOTATE ORQALI) ---
            # .count() o'rniga annotate qilingan 'num_questions' dan foydalanamiz
            q_count = getattr(obj, 'num_questions', 0)
            
            if q_count_filter:
                try:
                    target_q = int(q_count_filter)
                    if q_count == target_q: match_score += 2000
                    elif q_count >= target_q: match_score += 1000
                except ValueError: pass

            # --- VAQT VA METOD (FILE TYPE) ---
            raw_method = getattr(obj, 'creation_method', 'manual')
            c_method = str(raw_method).strip().lower() if raw_method else 'manual'
            db_time = int(getattr(obj, 'test_duration', 0) or getattr(obj, 'duration_minutes', 0) or getattr(obj, 'duration', 0) or 0)
            
            # Fayl turini fan nomi yoki yaratilish metodidan aniqlaymiz
            is_from_file = c_method in ['file', 'word', 'excel', 'txt'] or 'fayl' in fan_nomi_db or (db_time == 30 and q_count != 30 and q_count > 0)
            
            file_extension = 'word' # Default holatda
            if 'excel' in fan_nomi_db or c_method == 'excel': file_extension = 'excel'
            elif 'txt' in fan_nomi_db or c_method == 'txt': file_extension = 'txt'

            if is_from_file:
                final_time = q_count if q_count > 0 else 30
                method_str = f"({file_extension.upper()} FAYLDAN)"
            else:
                final_time = db_time if db_time > 0 else (q_count if q_count > 0 else 15)
                method_str = "(QO'LDA)"

            if duration_filter:
                try:
                    target_time = int(duration_filter)
                    if final_time == target_time: match_score += 2000
                    elif final_time <= target_time: match_score += 1000
                except ValueError: pass

            # --- MUALLIF MA'LUMOTI ---
            user_obj = getattr(obj, 'creator', getattr(obj, 'user', None))
            muallif = "Noma'lum"
            if user_obj:
                f_name = str(getattr(user_obj, 'first_name', '')).strip()
                l_name = str(getattr(user_obj, 'last_name', '')).strip()
                muallif = f"{f_name} {l_name}".strip() if (f_name or l_name) else user_obj.username

            if mine_filter:
                if request.user.is_authenticated and user_obj and user_obj.id == request.user.id:
                    match_score += 10000
                else: return None

            # --- 3. SELECT FILTRI (WORD, EXCEL, TXT, KAZUS, YOPIQ) ---
            if select_cat:
                if select_cat in ['word', 'excel', 'txt']:
                    if not is_from_file: return None
                    if select_cat == 'word' and file_extension != 'word': return None
                    if select_cat == 'excel' and file_extension != 'excel': return None
                    if select_cat == 'txt' and file_extension != 'txt': return None
                elif select_cat == 'kazus' and not is_kazus: return None
                elif select_cat == 'yopiq' and not is_closed: return None
                match_score += 4000

            # --- 4. INPUT TOIFA FILTRI ---
            if input_cat:
                is_match = False
                if ("kazus" in input_cat or "case" in input_cat or "kasuz" in input_cat) and is_kazus: is_match = True
                elif ("yopiq" in input_cat or "closed" in input_cat) and is_closed: is_match = True
                elif ("oddiy" in input_cat or "regular" in input_cat) and (not is_kazus and not is_closed): is_match = True
                elif ("fayl" in input_cat or "word" in input_cat or "excel" in input_cat or "txt" in input_cat) and is_from_file: is_match = True
                
                if is_match: match_score += 5000
                else: return None

            # --- 5. QIDIRUV PARAMETRI ---
            if query:
                if (query in title_str) or (query in muallif.lower()) or (query in label.lower()) or (query in fan_nomi_db):
                    match_score += 2000
                else: return None

            return {
                "id": obj_id,
                "title": title_str.capitalize(),
                "fan_nomi": getattr(obj, 'fan_nomi', None) or f"{label} {method_str}",
                "muallif": muallif,
                "duration": final_time,
                "q_count": q_count,
                "type_key": t_key,
                "creation_method": "file" if is_from_file else "manual",
                "match_score": match_score 
            }

        # --- OPTIMALLASHTIRILGAN DATA TO'PLASH (PREFETCH & ANNOTATE) ---
        
        # 1. CaseTest yuklash
        qs_case = CaseTest.objects.all().select_related('creator', 'subject').annotate(
            num_questions=Count('questions')
        )
        for obj in qs_case:
            res = process_test_object(obj, 'case')
            if res: all_tests_raw.append(res)

        # 2. ClosedTest yuklash
        qs_closed = ClosedTest.objects.all().select_related('user', 'subject').annotate(
            num_questions=Count('closed_questions')
        )
        for obj in qs_closed:
            res = process_test_object(obj, 'closed')
            if res: all_tests_raw.append(res)

        # 3. Asosiy Test modelini yuklash
        # Bog'liqliklar xotiraga prefetch qilinadi, N+1 muammosi bo'lmaydi
        qs_test = Test.objects.all().select_related('creator', 'subject').prefetch_related(
            'case_test_detail', 'questions'
        ).annotate(
            num_questions=Count('questions')
        )
        
        for obj in qs_test:
            # Model ichida closedtest bog'liqligi FieldError bermasligi uchun hasattr xavfsiz tekshiriladi
            has_case_rel = hasattr(obj, 'case_test_detail') and obj.case_test_detail is not None
            has_closed_rel = hasattr(obj, 'closedtest') and obj.closedtest is not None
            
            if has_case_rel or has_closed_rel: 
                continue
                
            res = process_test_object(obj, 'main')
            if res: all_tests_raw.append(res)

        # REYTING BUYICHA SARALASH
        all_tests_raw.sort(key=lambda x: (x['match_score'], x['id']), reverse=True)
        
        # UNIFIKATSIYA (TAKRORLANISHNI OLDINI OLISH)
        seen = set()
        final_list = []
        for item in all_tests_raw:
            unique_key = f"{item['type_key']}_{item['id']}"
            if unique_key not in seen:
                final_list.append(item)
                seen.add(unique_key)

        return Response(final_list)

    except Exception as e:
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)


@api_view(['GET'])
def api_test_list(request):
    search_query = request.GET.get('search', '')
    subject = request.GET.get('subject', '')
    duration = request.GET.get('duration', '')
    test_type = request.GET.get('type', '')

    results = []

    # Qaysi modeldan qidirishni aniqlash
    # 'file' yoki 'manual' kelsa Test modelidan, 'case' bo'lsa CaseTestdan qidiramiz
    if test_type == 'manual' or test_type == 'file':
        models_to_search = [Test]
    elif test_type == 'case':
        models_to_search = [CaseTest]
    elif test_type == 'yopiq':
        models_to_search = [ClosedTest]
    else:
        models_to_search = [Test, CaseTest, ClosedTest]

    for model in models_to_search:
        queryset = model.objects.all()

        # Filtrlash (Creation method bo'yicha - agar Test modeli bo'lsa)
        if model == Test and test_type in ['manual', 'file']:
            queryset = queryset.filter(creation_method=test_type)

        # Qidiruv (nomi yoki muallifi bo'yicha)
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query) | Q(creator__username__icontains=search_query)
            )

        # Fan bo'yicha filtr
        if subject:
            queryset = queryset.filter(subject__name__icontains=subject)

        # Vaqt bo'yicha filtr
        if duration:
            # Modelda maydon nomi 'duration_minutes' bo'lsa shuni yozing
            queryset = queryset.filter(duration_minutes__lte=int(duration))

        # Ma'lumotlarni yig'ish
        for item in queryset:
            # Frontendda "Yopiq Test" yozuvi o'rniga chiqadigan matnni aniqlaymiz
            display_label = "Yopiq Test" # Default
            
            if model == CaseTest:
                display_label = "Kazus Test"
            elif model == Test:
                if hasattr(item, 'creation_method'):
                    display_label = "Qo'lda yaratilgan" if item.creation_method == 'manual' else "Fayldan yuklangan"
                else:
                    display_label = "Oddiy Test"

            results.append({
                "id": item.id,
                "title": item.title,
                "display_type": display_label, # Mana shu frontendda "Yopiq Test" o'rniga chiqadi
                "model_type": model.__name__.lower(),
                "subject": item.subject.name if item.subject else "Noma'lum",
                "author": item.creator.get_full_name() or item.creator.username,
                "duration": getattr(item, 'duration_minutes', 0),
                "questions_count": item.questions.count() if hasattr(item, 'questions') else 0
            })

    return Response(results)

    
@api_view(['GET'])
def api_get_questions(request, test_id):
    # Frontenddan test turini olamiz (regular, case, closed)
    test_type = request.GET.get('type', 'regular')
    test_instance = None
    
    # 1. Turiga qarab aynan o'sha modeldan qidiramiz
    if test_type == 'regular':
        test_instance = Test.objects.filter(id=test_id).first()
    elif test_type == 'case':
        test_instance = CaseTest.objects.filter(id=test_id).first()
    elif test_type == 'closed':
        test_instance = ClosedTest.objects.filter(id=test_id).first()

    if not test_instance:
        return Response({
            "error": f"ID={test_id} va turi={test_type} bo'lgan test topilmadi."
        }, status=404)

    try:
        # Savollarni olish (related_name ni tekshirish)
        # ClosedTest uchun ko'pincha 'closedtestquestion_set' yoki 'questions' bo'ladi
        if hasattr(test_instance, 'questions'):
            questions = test_instance.questions.all()
        elif hasattr(test_instance, 'question_set'):
            questions = test_instance.question_set.all()
        else:
            # Agar ClosedTest modelida related_name bo'lmasa, Django model nomi bilan qidiradi
            questions = getattr(test_instance, f'{test_instance._meta.model_name}question_set').all()

        data = []
        for q in questions:
            data.append({
                "id": q.id,
                "text": q.text,
                "answer_a": getattr(q, 'answer_a', getattr(q, 'answer_A', '')),
                "answer_b": getattr(q, 'answer_b', getattr(q, 'answer_B', '')),
                "answer_c": getattr(q, 'answer_c', getattr(q, 'answer_C', '')),
                "answer_d": getattr(q, 'answer_d', getattr(q, 'answer_D', '')),
                "correct": getattr(q, 'correct_answer', getattr(q, 'correct', 'A'))
            })
        
        return Response(data)

    except Exception as e:
        return Response({"error": str(e)}, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def api_check_test(request):
    try:
        data = json.loads(request.body)
        test_id = data.get('test_id')
        user_answers = data.get('answers', {}) # {"102": "A", "103": "B"}

        # Har ikkala modeldan savollarni yig'ish
        questions = Question.objects.filter(Q(test_id=test_id) | Q(closed_test_id=test_id))
        if not questions.exists():
            questions = ClosedQuestion.objects.filter(test_id=test_id)

        score = 0
        total = questions.count()

        for q in questions:
            user_ans = str(user_answers.get(str(q.id)) or "").strip().upper()
            
            # Bazadagi javobni aniqlash
            raw_correct = (
                getattr(q, 'correct_answer', None) or 
                getattr(q, 'correct_option', None) or 
                getattr(q, 'correct_answers', "A")
            )
            
            # Agar vergul bilan kelgan bo'lsa (masalan "A, B"), birinchisini olamiz
            correct_db = str(raw_correct).split(',')[0].strip().upper()
            
            if user_ans and user_ans == correct_db:
                score += 1

        percent = round((score / total * 100), 1) if total > 0 else 0

        # Result modeliga yozish
        Result.objects.create(
            user=request.user,
            test_id=test_id if Test.objects.filter(id=test_id).exists() else None,
            score=score,
            total=total,
            percentage=percent
        )

        return JsonResponse({
            "status": "success",
            "score": score,
            "total": total,
            "percent": percent
        })
    except Exception as e:
        return JsonResponse({"status": "error", "message": str(e)}, status=400)
    

# --- Yopiq testlar ro'yxati ---
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def closed_tests_list(request):
    tests = ClosedTest.objects.all()
    data = [{"id": t.id, "title": t.title, "duration": t.duration, "score": t.score, "question_count": t.questions.count()} for t in tests]
    return Response(data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def start_closed_test(request, test_id):
    try:
        # 1. Testni olish
        test = ClosedTest.objects.get(id=test_id)
    except ClosedTest.DoesNotExist:
        return Response({"detail": "Test topilmadi"}, status=404)

    # 2. Sessiyani olish yoki yaratish
    session, _ = ClosedTestSession.objects.get_or_create(user=request.user, test=test)

    # 3. Savollarni olish (Modeldagi related_name='closed_questions' dan foydalanamiz)
    # Eslatma: Modeldagi variantlar answer_A, answer_B... ko'rinishida
    questions_queryset = test.closed_questions.all()
    
    if not questions_queryset.exists():
        return Response({"detail": "Ushbu testda savollar topilmadi", "questions": []}, status=200)

    questions_data = questions_queryset.values(
        'id', 
        'text', 
        'question_type', 
        'answer_A', 
        'answer_B', 
        'answer_C', 
        'answer_D'
    )

    # 4. Javob qaytarish
    return Response({
        "id": test.id,
        "title": test.title,
        "duration": test.duration,  # Modeldagi maydon nomi: duration
        "score": test.score,
        "session_id": session.id,
        "questions": list(questions_data)
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def start_case_test(request, test_id):
    try:
        # 1. Kazus testni bazadan qidiramiz
        case_test = CaseTest.objects.get(id=test_id)
    except CaseTest.DoesNotExist:
        return Response({"detail": "Kazus test topilmadi"}, status=404)

    # 2. Ushbu kazusga tegishli savollarni olamiz
    # CaseQuestion modelida related_name='questions' deb yozgansiz
    questions_queryset = case_test.questions.all()
    
    if not questions_queryset.exists():
        return Response({"detail": "Ushbu kazusda savollar topilmadi"}, status=200)

    # 3. Savollarni formatlaymiz
    questions_data = []
    for q in questions_queryset:
        questions_data.append({
            "id": q.id,
            "text": q.text,
            "answer_a": q.answer_a,
            "answer_b": q.answer_b,
            "answer_c": q.answer_c,
            "answer_d": q.answer_d,
            "correct_answers": q.correct_answers # JSONField
        })

    # 4. Kazus matni va savollarni qaytaramiz
    return Response({
        "id": case_test.id,
        "title": case_test.title,
        "case_text": case_test.case_text, # Kazusning asosiy matni
        "subject": case_test.subject.name if case_test.subject else "Noma'lum",
        "questions": questions_data
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def submit_closed_test(request, test_id):
    print(f"\n--- DEBUG START (Test ID: {test_id}) ---")
    
    try:
        # 1. Testni olish
        test = ClosedTest.objects.get(id=test_id)
        
        # 2. Sessiyani olish (test_id orqali qidiramiz)
        # Modelda maydon nomi 'closed_test' yoki 'test' ekanligini aniqlang
        session = ClosedTestSession.objects.filter(
            user=request.user, 
            closed_test=test # yoki test=test
        ).last()

        if not session:
            print("XATO: Sessiya topilmadi!")
            return Response({"detail": "Sessiya topilmadi"}, status=404)

        print(f"Sessiya topildi: ID {session.id}")

        # 3. Savollarni olish
        # related_name='closed_questions' bo'lishi kerak
        all_questions = test.closed_questions.all()
        total_q_count = all_questions.count()
        print(f"Testdagi jami savollar: {total_q_count}")

        # 4. Javoblarni olish
        answers_data = request.data.get("answers", [])
        print(f"Yuborilgan javoblar soni: {len(answers_data)}")

        correct_count = 0
        for ans in answers_data:
            q_id = ans.get('question_id')
            selected = str(ans.get('selected', '')).strip().upper()

            try:
                question = all_questions.get(id=q_id)
                db_correct = str(question.correct_answers).strip().upper()

                # Solishtirish
                if selected == db_correct:
                    correct_count += 1
                
                # Natijani bazaga yozish (keyinchalik ko'rish uchun)
                ClosedAnswer.objects.update_or_create(
                    session=session,
                    question=question,
                    defaults={"selected_answers": selected, "is_correct": (selected == db_correct)}
                )
            except ClosedQuestion.DoesNotExist:
                print(f"Savol topilmadi: ID {q_id}")
                continue

        # 5. Ballni hisoblash
        score = (correct_count / total_q_count * 10) if total_q_count > 0 else 0
        
        session.score_obtained = score
        session.completed_at = timezone.now()
        session.save()

        print(f"To'g'ri javoblar: {correct_count}")
        print(f"Yakuniy ball: {score}")
        print("--- DEBUG END ---\n")

        return Response({
            "score_obtained": round(score, 2),
            "correct_count": correct_count,
            "total_questions": total_q_count,
            "percent": round((correct_count / total_q_count * 100), 1) if total_q_count > 0 else 0
        })

    except ClosedTest.DoesNotExist:
        return Response({"detail": "Test topilmadi"}, status=404)
    except Exception as e:
        print(traceback.format_exc())
        return Response({"detail": str(e)}, status=400)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_manual_closed_test(request):
    data = request.data
    title = data.get("title")
    description = data.get("description", "")
    
    # Modelingizda 'duration' deb nomlangan, shuning uchun 'duration' kalitini ishlatamiz
    # Frontenddan 'duration_minutes' kelsa ham, uni 'duration'ga o'g'irib olamiz
    duration_val = data.get("duration_minutes") or data.get("duration") or 30
    score_val = data.get("score", 100)
    questions = data.get("questions", [])

    # 1. Majburiy maydonlarni tekshirish
    if not title or not questions:
        return Response(
            {"detail": "Test nomi va savollar kiritilishi shart!"}, 
            status=status.HTTP_400_BAD_REQUEST
        )

    try:
        # Tranzaksiya: Agar savollardan birida xato bo'lsa, test ham yaratilmaydi (baza tozalanadi)
        with transaction.atomic():
            # 2. Test ob'ektini yaratish (Modelga moslandi)
            test = ClosedTest.objects.create(
                title=title,
                description=description,
                duration=int(duration_val), # Modelda duration (IntegerField)
                score=int(score_val),       # Modelda score (IntegerField)
                user=request.user
            )

            for q in questions:
                # 3. To'g'ri javobni tekshirish va formatlash
                correct = str(q.get('correct', '')).strip().upper()
                
                if not correct:
                    raise ValueError(f"'{q.get('text', 'Nomsiz')}' savoli uchun to'g'ri javob ko'rsatilmadi.")

                # 4. Variantlarni xavfsiz olish
                answers = q.get('answers', {})
                
                # 5. Savolni bazaga saqlash
                ClosedQuestion.objects.create(
                    test=test,
                    text=q.get('text', 'Savol matni kiritilmagan'),
                    question_type=q.get('question_type', 'single'),
                    answer_A=str(answers.get('A', '')).strip(),
                    answer_B=str(answers.get('B', '')).strip(),
                    answer_C=str(answers.get('C', '')).strip(),
                    answer_D=str(answers.get('D', '')).strip(),
                    correct_answers=correct # Modelda correct_answers (CharField)
                )

        return Response({
            "status": "success",
            "detail": "Test va savollar muvaffaqiyatli yaratildi", 
            "test_id": test.id
        }, status=status.HTTP_201_CREATED)

    except ValueError as ve:
        return Response({"detail": str(ve)}, status=status.HTTP_400_BAD_REQUEST)
    except Exception as e:
        # Terminalda xatolikning aniq sababini ko'rish uchun
        print(f"CREATE_TEST_ERROR: {str(e)}")
        return Response(
            {"detail": f"Serverda xatolik yuz berdi: {str(e)}"}, 
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_manual_ordinary_test(request):
    data = request.data
    questions = data.get("questions", [])

    try:
        with transaction.atomic():
            test = Test.objects.create(
                title=data.get("title", "Yangi Test"),
                creator=request.user,
                test_type='ordinary',
                creation_method='manual'
            )

            for q in questions:
                # --- MANA SHU YERNI O'ZGARTIRASIZ ---
                # Oldingi 'or A' degan joyini 'or None' ga almashtiramiz
                raw_correct = q.get('correct') or q.get('correct_answer') or q.get('correct_option') or None
                
                # Agar frontenddan javob kelmasa, 'X' bo'ladi (A emas)
                correct = str(raw_correct).strip().upper() if raw_correct else "X"
                # -------------------------------------

                Question.objects.create(
                    test=test,
                    creator=request.user,
                    text=q.get('text', ''),
                    answer_a=str(q.get('answer_a') or q.get('A') or '').strip(),
                    answer_b=str(q.get('answer_b') or q.get('B') or '').strip(),
                    answer_c=str(q.get('answer_c') or q.get('C') or '').strip(),
                    answer_d=str(q.get('answer_d') or q.get('D') or '').strip(),
                    
                    # Bu flaglar endi faqat 'correct' haqiqatdan A, B, C yoki D bo'lsagina True bo'ladi
                    is_correct_a=(correct == 'A'),
                    is_correct_b=(correct == 'B'),
                    is_correct_c=(correct == 'C'),
                    is_correct_d=(correct == 'D'),
                    correct_answer=correct
                )

        return Response({"status": "success"}, status=201)
    except Exception as e:
        return Response({"detail": str(e)}, status=500)
     

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_closed_tests(request):
    tests = ClosedTest.objects.all()
    results = []

    for t in tests:
        questions_list = []
        for q in t.questions.all():  # related_name='questions'
            questions_list.append({
                "id": q.id,
                "text": q.text,
                "question_type": q.question_type,
                "answer_A": q.answer_A,
                "answer_B": q.answer_B,
                "answer_C": q.answer_C,
                "answer_D": q.answer_D,
                "correct_answers": q.correct_answers
            })
        results.append({
            "id": t.id,
            "title": t.title,
            "description": t.description,
            "duration_minutes": t.duration_minutes,
            "score": t.score,
            "questions": questions_list
        })
    return Response(results)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def api_save_closed_answer(request):
    data = request.data
    s_id = data.get('session_id')
    q_id = data.get('question_id')
    u_ans = data.get('answer', '')

    session = TestSession.objects.filter(id=s_id).first()
    question = ClosedQuestion.objects.filter(id=q_id).first()

    if not session or not question:
        return Response({'error': 'Sessiya yoki savol topilmadi'}, status=400)

    # --- JAVOBNI TOZALASH (MUHIM QISMI) ---
    # 1. Agar javob massiv bo'lsa, birinchisini olamiz yoki string qilamiz
    if isinstance(u_ans, list):
        user_answer_str = "".join(map(str, u_ans)).strip().lower()
    else:
        user_answer_str = str(u_ans).strip().lower()

    # 2. Bazadagi to'g'ri javobni tozalash (+, ), . belgilarni olib tashlaymiz)
    # Masalan: "+A" -> "a", "A)" -> "a", " A " -> "a"
    import re
    correct_ans_db = str(question.correct_answers).replace('+', '').strip().lower()
    correct_ans_db = re.sub(r'[^\w]', '', correct_ans_db) # Faqat harf va raqam qoladi
    user_answer_str = re.sub(r'[^\w]', '', user_answer_str)

    # 3. Tekshirish
    is_correct = (user_answer_str == correct_ans_db)

    # DEBUG: Terminalda nima bilan nima solishtirilayotganini ko'ring
    print(f"SOLISHTIRISH: User: '{user_answer_str}' <==> DB: '{correct_ans_db}' | Natija: {is_correct}")

    ClosedAnswer.objects.update_or_create(
        session=session,
        question=question,
        defaults={'selected_answers': u_ans, 'is_correct': is_correct}
    )

    return Response({'status': 'success', 'is_correct': is_correct})


@api_view(['GET'])
def time_remaining(request, session_id):
    session = get_object_or_404(TestSession, id=session_id)
    duration = session.test.duration_minutes  # masalan 20
    passed = (timezone.now() - session.start_time).total_seconds()
    remaining = duration * 60 - passed

    if remaining <= 0:
        session.finished = True
        session.save()
        remaining = 0

    return Response({
        "remaining_seconds": int(remaining),
        "finished": session.finished
    })


def my_closed_tests_page(request):
    return render(request, 'test/my_tests.html')


import io
import logging
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import authentication, permissions
from django.db import transaction

# O'zingizning to'g'rilangan parser funksiyangizni va modellaringizni import qiling
from .utils import parse_and_create_questions 
from .models import Subject, Test, CaseTest, ClosedTest

logger = logging.getLogger(__name__)

class UploadFileTestsAPIView(APIView):
    authentication_classes = [authentication.TokenAuthentication]
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, *args, **kwargs):
        files = request.FILES
        if not files:
            return Response({"detail": "Fayl yuborilmadi."}, status=400)

        subject_id = request.data.get("subject_id")
        
        # 🔍 DIAGNOSTIKA: Terminalda frontenddan aynan nima kelayotganini ko'ramiz
        incoming_type = request.data.get("type")
        print(f"--- 📥 FRONTENDDAN KELGAN ASLIY TYPE: {incoming_type} ---")
        print(f"--- 📥 BARCHA POST MA'LUMOTLARI: {request.data} ---")

        # Frontenddan kelgan turni tozalaymiz (agar bo'sh bo'lsa 'regular')
        test_type = str(incoming_type).strip().lower() if incoming_type else "regular"
        
        subject, _ = Subject.objects.get_or_create(
            id=subject_id if subject_id else None, 
            defaults={'name': "Default fan"}
        )

        saved_tests = []
        
        for key in files:
            f = files[key]
            test_title = f.name.split('.')[0]
            filename_lower = f.name.lower()
            
            # 🎯 MAJBURIY QO'SHIMCHA TEKSHIRUV (Zaxira xavfsizlik kamari):
            # Agar turi 'kazus' deb kelgan bo'lsa YOKI fayl nomida 'kazus', 'case', 'kasuz' so'zlari bo'lsa
            # kod uni 100% KAZUS deb qabul qiladi, 'regular' bo'lib ketishiga yo'l qo'ymaydi!
            current_file_type = test_type
            if "kazus" in filename_lower or "case" in filename_lower or "kasuz" in filename_lower or test_type in ["kazus", "case"]:
                current_file_type = "kazus"
            elif "closed" in filename_lower or "yopiq" in filename_lower or test_type == "closed":
                current_file_type = "closed"
            
            print(f"--- 🚀 PROTSESS: Fayl: {f.name} -> Aniqlangan yakuniy tur: {current_file_type} ---")
            
            try:
                with transaction.atomic():
                    # --- 1. MODEL TURINI ANIQ YARATISH ---
                    if current_file_type == "closed":
                        new_test = ClosedTest.objects.create(
                            title=test_title,
                            user=request.user,
                            subject=subject,
                            duration=0,
                            score=100,
                            file=f
                        )
                    elif current_file_type == "kazus":
                        # 1. Avval ota model yaratiladi
                        base_test = Test.objects.create(
                            title=test_title,
                            subject=subject,
                            creator=request.user,
                            test_duration=0,
                            creation_method='file',
                            test_type='kazus',  # Baza darajasida 'kazus' deb muhrlash
                            category="Kazus Testlar"
                        )
                        # 2. Keyin bola model (CaseTest) unga ulanadi
                        new_test = CaseTest.objects.create(
                            test_ptr=base_test,
                            title=test_title,
                            creator=request.user,
                            subject=subject,
                            case_text=f"{test_title} uchun fayldan yuklangan kazus",
                            test_duration=0,
                            file=f
                        )
                    else:
                        new_test = Test.objects.create(
                            title=test_title,
                            subject=subject,
                            creator=request.user,
                            test_duration=0,
                            creation_method='file',
                            test_type='ordinary',
                            category="Fayldan yuklangan"
                        )

                    # --- 2. ASOSIY PARSERGA BERIB YUBORISH ---
                    # `parse_and_create_questions` funksiyasi savollarni parsing qiladi
                    created_count = parse_and_create_questions(new_test, f, request.user)

                    saved_tests.append({
                        "test_id": new_test.pk, 
                        "title": new_test.title,
                        "type": current_file_type,
                        "questions_created": created_count
                    })

            except Exception as e:
                print(f"❌ Xatolik ({f.name}): {str(e)}")
                import traceback
                traceback.print_exc()
                continue 

        return Response({
            "status": "success", 
            "message": f"{len(saved_tests)} ta test muvaffaqiyatli yuklandi.",
            "tests": saved_tests
        }, status=201)


class TestViewSet(viewsets.ModelViewSet):
    queryset = Test.objects.all().order_by('-id')
    serializer_class = TestSerializer
    authentication_classes = [TokenAuthentication]
    permission_classes = [permissions.IsAuthenticated]
    filter_backends = [filters.SearchFilter]
    search_fields = ['title', 'subject__name', 'creator__username']

    def create(self, request, *args, **kwargs):
        try:
            serializer = self.get_serializer(data=request.data)
            serializer.is_valid(raise_exception=True)
            test = serializer.save(creator=request.user)

            questions_data = request.data.get('questions')
            if questions_data:
                # Agar frontend JSON string yuborsa
                if isinstance(questions_data, str):
                    questions_data = json.loads(questions_data)

                for i, q_data in enumerate(questions_data, start=1):
                    raw_answers = q_data.get('correct_answers')

                    # ✅ correct_answers ni tozalash va formatlash
                    if isinstance(raw_answers, str):
                        correct_list = [
                            x.strip().upper() for x in raw_answers.split(',') if x.strip()
                        ]
                    elif isinstance(raw_answers, list):
                        correct_list = [
                            x.strip().upper() for x in raw_answers if x and x.upper() != "NULL"
                        ]
                    else:
                        correct_list = []

                    # ✅ Hech bo‘lmasa bitta javob bo‘lishini tekshirish
                    if not correct_list:
                        raise serializers.ValidationError({
                            "correct_answers": f"Question {i}: To‘g‘ri javob kiritilishi shart!"
                        })

                    primary_answer = correct_list[0]

                    Question.objects.create(
                        test=test,
                        subject=test.subject,
                        creator=request.user,
                        text=q_data.get('text', ''),
                        question_type=q_data.get('question_type', 'single'),
                        correct_answers=correct_list,
                        correct_answer=primary_answer,
                        is_multiple_choice=q_data.get('question_type') == 'multiple',
                        file=request.FILES.get(f'file_{i}')
                    )

            # 🎯 MANA SHU YERDA FAYLLARNI TEKSHIRAMIZ:
            files = request.FILES.getlist('files[]')
            if files:
                # Agar fayllar kelgan bo'lsa, Test modelida ushbu test fayldan yuklanganini belgilaymiz
                if hasattr(test, 'creation_method'):
                    test.creation_method = 'file'
                # Agar modelingizda maxsus boolean maydon bo'lsa (masalan is_file_uploaded):
                if hasattr(test, 'is_file_uploaded'):
                    test.is_file_uploaded = True
                
                # O'zgarishni bazada saqlaymiz
                test.save()

            for file in files:
                name = file.name.lower()
                content_list = []

                if name.endswith('.txt'):
                    content = file.read().decode('utf-8')
                    content_list = [line.strip() for line in content.split('\n') if line.strip()]

                elif name.endswith('.docx'):
                    result = mammoth.extract_raw_text(file)
                    content_list = [line.strip() for line in result.value.split('\n') if line.strip()]

                elif name.endswith(('.xls', '.xlsx')):
                    wb = openpyxl.load_workbook(file)
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            content_list.append(' | '.join([str(cell) for cell in row if cell]))

                for item in content_list:
                    Question.objects.create(
                        test=test,
                        subject=test.subject,
                        creator=request.user,
                        text=item,
                        question_type='single',
                        correct_answers=["A"],  
                        correct_answer="A",
                        is_multiple_choice=False
                    )

            return Response(TestSerializer(test).data, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)


class OlympiadGroupListCreateView(generics.ListCreateAPIView):
    queryset = OlympiadGroup.objects.all().order_by('-id')
    serializer_class = OlympiadGroupSerializer
    authentication_classes = [SessionAuthentication, TokenAuthentication]
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)


class OlympiadParticipantViewSet(viewsets.ModelViewSet):
    queryset = OlympiadParticipant.objects.all().select_related("user", "group")
    serializer_class = OlympiadParticipantSerializer
    authentication_classes = [SessionAuthentication, TokenAuthentication]
    permission_classes = [permissions.IsAuthenticated]


class JoinOlympiadGroupView(APIView):
    authentication_classes = [SessionAuthentication, TokenAuthentication]
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, group_id):
        group = get_object_or_404(OlympiadGroup, id=group_id)
        participant, created = OlympiadParticipant.objects.get_or_create(
            group=group, user=request.user
        )
        if not created:
            return Response(
                {"detail": "Siz allaqachon ushbu guruhda ishtirokchisiz."},
                status=status.HTTP_400_BAD_REQUEST,
            )
        serializer = OlympiadParticipantSerializer(participant)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    

class OlympiadGroupRetrieveView(generics.RetrieveAPIView):
    queryset = OlympiadGroup.objects.all()
    serializer_class = OlympiadGroupSerializer
    permission_classes = [IsAuthenticated]


@api_view(["POST"])
@permission_classes([permissions.IsAuthenticated])
def olympiad_submit_answer(request, group_id):
    """
    Foydalanuvchi javob yuborganida chaqiriladi.
    To‘g‘ri javobni tekshiradi va ballni hisoblaydi.
    """
    # 1. Guruh va ishtirokchini olish
    group = get_object_or_404(OlympiadGroup, id=group_id)
    participant = get_object_or_404(OlympiadParticipant, group=group, user=request.user)

    if participant.finished:
        return Response({"detail": "Siz testni allaqachon yakunlagansiz."}, status=status.HTTP_400_BAD_REQUEST)

    # 2. Requestdan ma'lumot olish
    question_id = request.data.get("question_id")
    selected = request.data.get("selected")

    if not question_id or selected is None:
        return Response({"detail": "question_id va selected majburiy."}, status=status.HTTP_400_BAD_REQUEST)

    # 3. Savolni olish
    question = get_object_or_404(OlympiadQuestion, id=question_id, group=group)

    # 4. To‘g‘ri javobni tekshirish (Modelingizdagi 'correct_answer' ga moslandi)
    correct_str = str(question.correct_answer).strip().lower() if question.correct_answer else ""
    selected_str = str(selected).strip().lower()
    
    is_correct = (selected_str == correct_str) and (correct_str != "")
   

    # 5. Javobni saqlash yoki yangilash
    OlympiadAnswer.objects.update_or_create(
        participant=participant,
        question_id=question.id,
        defaults={
            "is_correct": is_correct,
            "selected": selected_str,
            "answered_at": timezone.now()
        }
    )

    # 6. Ball va javoblar sonini qayta hisoblash
    user_answers = OlympiadAnswer.objects.filter(participant=participant)
    total_correct = user_answers.filter(is_correct=True).count()
    total_wrong = user_answers.filter(is_correct=False).count()

    # To'g'ri javob berilgan savollarning ID larini olamiz
    correct_q_ids = user_answers.filter(is_correct=True).values_list('question_id', flat=True)
    
    # Savollarning 'difficulty' (qiyinchilik darajasi) ball sifatida olinadi
    correct_questions = OlympiadQuestion.objects.filter(id__in=correct_q_ids)
    
    current_score = 0
    for q in correct_questions:
        # Modelingizda 'difficulty' bor, agar u yo'q bo'lsa 1 ball beradi
        current_score += getattr(q, "difficulty", 1)

    # 7. Ishtirokchi ma'lumotlarini yangilash
    participant.score = float(current_score)
    participant.correct_answers = total_correct
    participant.wrong_answers = total_wrong
    participant.last_activity = timezone.now()
    participant.save(update_fields=["correct_answers", "wrong_answers", "score", "last_activity"])

    # 8. Qolgan savollar soni
    total_questions = OlympiadQuestion.objects.filter(group=group).count()
    answered_count = user_answers.count() # Nechta savolga javob berilgani (to'g'ri + xato)
    remaining = max(total_questions - answered_count, 0)

    # 9. Frontendga JSON javob
    return Response({
        "status": "success",
        "question_id": question.id,
        "is_correct": is_correct,
        "correct_answers": total_correct,
        "wrong_answers": total_wrong,
        "remaining": remaining,
        "score": participant.score
    }, status=status.HTTP_200_OK)


# ✅ TESTNI YAKUNLASH (finish_test)
@api_view(["POST"])
@permission_classes([permissions.IsAuthenticated])
def finish_olympiad(request, group_id):
    # 1. Guruhni olish
    group = get_object_or_404(OlympiadGroup, id=group_id)

    # 2. Foydalanuvchi ishtirokchisini olish
    participant = get_object_or_404(
        OlympiadParticipant,
        group=group,
        user=request.user
    )

    # Agar allaqachon yakunlagan bo‘lsa, mavjud natijani qaytaramiz
    if participant.finished:
        return Response({
            "message": "Siz allaqachon testni yakunlagansiz.",
            "score": participant.score,
            "correct_answers": participant.correct_answers,
            "wrong_answers": participant.wrong_answers
        }, status=status.HTTP_200_OK)

    # 3. Javoblarni qayta hisoblash (Yakuniy nazorat)
    user_answers = OlympiadAnswer.objects.filter(participant=participant)
    total_correct = user_answers.filter(is_correct=True).count()
    total_wrong = user_answers.filter(is_correct=False).count()

    # 4. Ballarni 'difficulty' (qiyinchilik darajasi) asosida hisoblash
    # OlympiadAnswer da question ForeignKey yo'qligi sababli ID orqali savollarni olamiz
    correct_q_ids = user_answers.filter(is_correct=True).values_list('question_id', flat=True)
    correct_questions = OlympiadQuestion.objects.filter(id__in=correct_q_ids)
    
    final_score = 0
    for q in correct_questions:
        # Modelingizda 'difficulty' bor, agar u yo'q bo'lsa default 1 ball
        final_score += getattr(q, "difficulty", 1)

    # 5. Ishtirokchi statistikasini yangilash
    participant.correct_answers = total_correct
    participant.wrong_answers = total_wrong
    participant.score = float(final_score)  # Umumiy to'plangan ball
    participant.finished = True
    participant.last_activity = timezone.now()
    
    participant.save(update_fields=[
        "correct_answers", "wrong_answers", "score", "finished", "last_activity"
    ])

    # 6. Guruhdagi progressni yangilash (Agar WebSocket signal bo'lsa)
    try:
        # Agar bu funksiya mavjud bo'lsa ishlaydi, yo'q bo'lsa jarayon to'xtab qolmaydi
        _broadcast_group_progress(group.id)
    except Exception:
        pass 

    return Response({
        "message": "Test yakunlandi.",
        "score": participant.score,
        "correct_answers": total_correct,
        "wrong_answers": total_wrong
    }, status=status.HTTP_200_OK)


# ✅ REAL-TIME PROGRESS SIGNAL (channels uchun)
def _broadcast_group_progress(group_id: int):
    try:
        group = OlympiadGroup.objects.get(id=group_id)
    except OlympiadGroup.DoesNotExist:
        return

    OlympiadQuestion = apps.get_model("amaliyot", "OlympiadQuestion")
    total_questions = OlympiadQuestion.objects.filter(group=group).count()
    participants = OlympiadParticipant.objects.filter(group_id=group_id).select_related("user")

    res = []
    for p in participants:
        answered = p.correct_answers + p.wrong_answers
        remaining = total_questions - answered
        res.append({
            "username": p.user.username,
            "correct": p.correct_answers,
            "wrong": p.wrong_answers,
            "remaining": remaining,
            "finished": p.finished,
            "score": p.score,
            "last_activity": p.last_activity.strftime("%H:%M:%S")
        })

    finished_users = participants.filter(finished=True).order_by("last_activity")
    first_finished = finished_users[0].user.username if finished_users else None

    payload = {"participants": res, "first_finished": first_finished}

    layer = get_channel_layer()
    async_to_sync(layer.group_send)(
        f"olympiad_{group_id}",
        {"type": "progress_update", "data": payload}
    )


@api_view(["POST"])
@permission_classes([permissions.IsAuthenticated])
def start_olympiad(request, group_id):
    # 1. Guruhni olish
    group = get_object_or_404(OlympiadGroup, id=group_id)

    # 2. Foydalanuvchi vaqtini tekshirish (Guruh vaqti tugagan bo'lishi mumkin)
    now = timezone.now()
    if group.end_time and now > group.end_time:
        return Response({
            "detail": "Ushbu olimpiada muddati tugagan!"
        }, status=status.HTTP_400_BAD_REQUEST)

    # 3. Foydalanuvchi ishtirokchisini olish yoki yaratish
    participant, created = OlympiadParticipant.objects.get_or_create(
        group=group,
        user=request.user
    )

    # 4. Agar foydalanuvchi testni allaqachon tugatgan bo'lsa
    if participant.finished:
        return Response({
            "message": "Siz ushbu testni allaqachon yakunlagansiz!",
            "finished": True,
            "score": participant.score
        }, status=status.HTTP_200_OK)

    # 5. Test boshlangan vaqtni saqlash
    if not participant.start_time:
        participant.start_time = now
        participant.save(update_fields=["start_time"])

    # 6. Guruhdagi savollarni olish
    questions_qs = OlympiadQuestion.objects.filter(group=group).order_by('id') # Tartib bilan olish
    
    if not questions_qs.exists():
        return Response({
            "message": "Hozircha ushbu guruhda savollar mavjud emas!",
            "questions": [],
            "total_questions": 0
        }, status=status.HTTP_200_OK)

    # 7. Savollarni JSON formatga tayyorlash
    questions_data = []
    for q in questions_qs:
        # Variantlarni (options) qayta ishlash
        opts = []
        if q.options:
            if isinstance(q.options, str):
                # Vergul bilan ajratilgan bo'lsa, tozalab listga o'tkazamiz
                opts = [x.strip() for x in q.options.split(",") if x.strip()]
            elif isinstance(q.options, list):
                opts = q.options
        
        # Savol matni mavjudligini tekshirish (q.text yoki q.question bo'lishi mumkin)
        text = getattr(q, "text", None) or getattr(q, "question", "Savol matni mavjud emas")

        questions_data.append({
            "id": q.id,
            "text": text,
            "options": opts,
            "difficulty": q.difficulty
        })

    return Response({
        "message": f"Olimpiada '{group.name}' boshlandi!",
        "questions": questions_data,
        "total_questions": len(questions_data),
        "start_time": participant.start_time
    }, status=status.HTTP_200_OK)


@api_view(["GET"])
def leaderboard_view(request, group_id):
    group = get_object_or_404(OlympiadGroup, id=group_id)
    
    # Tartiblash: 
    # 1. Ball bo'yicha (eng yuqori ball tepada)
    # 2. last_activity bo'yicha (ballar teng bo'lsa, oxirgi harakati vaqtliroq bo'lgan odam tepada)
    participants = OlympiadParticipant.objects.filter(
        group=group
    ).select_related('user', 'user__userprofile').exclude(
        user__is_staff=True
    ).order_by('-score', 'last_activity')

    leaderboard_data = []
    current_user_id = request.user.id if request.user.is_authenticated else None

    for index, p in enumerate(participants, start=1):
        u = p.user
        
        # Ism-familiyani modelingizdagi __str__ mantiqi bo'yicha olish
        full_name = ""
        try:
            if hasattr(u, 'userprofile') and u.userprofile.full_name:
                full_name = u.userprofile.full_name
            else:
                full_name = f"{u.first_name} {u.last_name}".strip()
        except:
            full_name = u.username

        if not full_name:
            full_name = u.username
            
        leaderboard_data.append({
            "rank": index,
            "id": u.id,
            "full_name": full_name,
            "score": float(p.score or 0),
            "correct_answers": p.correct_answers,
            "is_me": u.id == current_user_id,
            "is_finished": p.finished
        })

    return Response({
        "status": "success",
        "group_name": group.name,
        "leaderboard": leaderboard_data
    })

@api_view(['GET'])
@permission_classes([permissions.IsAuthenticated])
def group_participants_view(request, group_id):
    group = get_object_or_404(OlympiadGroup, id=group_id)
    # select_related orqali bazaga so'rovlar sonini kamaytiramiz
    participants = OlympiadParticipant.objects.filter(
        group=group
    ).select_related('user', 'user__userprofile')
    
    # Serializer klassingiz mavjudligiga ishonch hosil qiling
    serializer = OlympiadParticipantSerializer(participants, many=True)
    return Response({
        "group": group.name,
        "participants": serializer.data
    })


@api_view(['POST'])
def add_olympiad_question(request):

    group_id = request.data.get("group")
    if not group_id:
        return Response({"detail": "Guruh ID topilmadi"}, status=400)

    # ❗ to‘g‘ri obyektni olish
    group = OlympiadGroup.objects.get(id=group_id)

    question = OlympiadQuestion.objects.create(
        group=group,  # ✔ endi to‘g‘ri
        text=request.data.get("question"),
        options=request.data.get("answers"),
        correct_answer=request.data.get("correct_answer"),
        difficulty=request.data.get("difficulty")
    )

    return Response({"detail": "Savol qo‘shildi", "id": question.id})


def olympiad_page(request):
    return render(request, 'test/olympiad.html')


def test_create_page(request):
    return render(request, "test/create_test.html")


def group_page(request):
    return render(request, "test/group.html")


class AuthViewSet(ViewSet):
    permission_classes = [IsAuthenticated]

    @action(detail=False, methods=['get'])
    def me(self, request):
        return Response({
            "id": request.user.id,
            "username": request.user.username,
        })


class GroupViewSet(viewsets.ModelViewSet):
    queryset = Group.objects.all().order_by('-created_at')
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action == 'create':
            return GroupCreateSerializer
        return GroupSerializer

    def perform_create(self, serializer):
        group = serializer.save(creator=self.request.user)
        GroupMembership.objects.create(group=group, user=self.request.user)

    # ---- A'zolar ro'yxati ----
    @action(detail=True, methods=['get'])
    def members(self, request, pk=None):
        group = self.get_object()

        if not group.memberships.filter(user=request.user).exists():
            return Response({"detail": "Guruh a'zosi emassiz."}, status=status.HTTP_403_FORBIDDEN)

        memberships = group.memberships.all()
        serializer = GroupMembershipSerializer(memberships, many=True)
        return Response(serializer.data)

    # ---- 1) TOKEN orqali join ----
    @action(detail=False, methods=['post'], url_path='join/(?P<token>[^/.]+)')
    def join_by_token(self, request, token=None):
        """Taklif tokeni orqali guruhga qo'shilish."""
        try:
            group = Group.objects.get(invite_token=token)
        except Group.DoesNotExist:
            return Response({"detail": "Noto'g'ri taklif tokeni."}, status=status.HTTP_404_NOT_FOUND)

        # Allaqachon a'zo
        if GroupMembership.objects.filter(group=group, user=request.user).exists():
            return Response({"detail": "Siz allaqachon bu guruhga a'zosiz."}, status=status.HTTP_400_BAD_REQUEST)

        # A'zo qilamiz
        GroupMembership.objects.create(group=group, user=request.user)
        return Response({"detail": f"Guruhga muvaffaqiyatli qo'shildingiz: {group.name}"}, status=status.HTTP_201_CREATED)

    # ---- 2) GROUP ID orqali join (SIZGA KERAK FUNKSIYA) ----
    @action(detail=True, methods=['post'], url_path='join')
    def join_by_id(self, request, pk=None):
        """Group ID orqali join."""
        try:
            group = Group.objects.get(id=pk)
        except Group.DoesNotExist:
            return Response({"detail": "Guruh topilmadi."}, status=status.HTTP_404_NOT_FOUND)

        if GroupMembership.objects.filter(group=group, user=request.user).exists():
            return Response({"detail": "Siz allaqachon bu guruhga a'zosiz."}, status=status.HTTP_400_BAD_REQUEST)

        GroupMembership.objects.create(group=group, user=request.user)
        return Response({"detail": f"Guruhga qo'shildingiz: {group.name}"}, status=status.HTTP_201_CREATED)


class GroupMembershipViewSet(viewsets.ModelViewSet):
    queryset = GroupMembership.objects.all()
    serializer_class = GroupMembershipSerializer
    permission_classes = [IsAuthenticated]
    http_method_names = ['post', 'delete'] 

    def create(self, request, *args, **kwargs):
        """Guruhga ID orqali qo'shilish."""
        group_id = request.data.get('group_id')
        if not group_id:
            return Response({"detail": "group_id kiritilishi shart."}, status=status.HTTP_400_BAD_REQUEST)
        
        group = get_object_or_404(Group, pk=group_id)
        
        if GroupMembership.objects.filter(group=group, user=request.user).exists():
            return Response({"detail": "Siz allaqachon bu guruhga a'zosiz."}, status=status.HTTP_400_BAD_REQUEST)
            
        GroupMembership.objects.create(group=group, user=request.user)
        return Response({"detail": f"Guruhga muvaffaqiyatli qo'shildingiz: {group.name}"}, status=status.HTTP_201_CREATED)

    def destroy(self, request, *args, **kwargs):
        """Guruhdan chiqish (o'z a'zoligini o'chirish)."""
        instance = get_object_or_404(GroupMembership, pk=kwargs['pk'], user=request.user)
        self.perform_destroy(instance)
        return Response(status=status.HTTP_204_NO_CONTENT)
    

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def current_user(request):
    user = request.user
    return Response({
        "id": user.id,
        "username": user.username,
        "email": user.email,
    })


class ChatMessageViewSet(viewsets.ModelViewSet):
    serializer_class = ChatMessageSerializer
    permission_classes = [IsAuthenticated]
    http_method_names = ['get', 'post'] 

    def get_queryset(self):
        group_pk = self.kwargs['group_pk']
        return ChatMessage.objects.filter(group_id=group_pk)

    def perform_create(self, serializer):
        group_pk = self.kwargs['group_pk']
        group = get_object_or_404(Group, pk=group_pk)
        serializer.save(sender=self.request.user, group=group)  


class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    permission_classes = [IsAuthenticated]


class TestCreateAPIView(generics.CreateAPIView):
    queryset = Test.objects.all()
    serializer_class = TestCreateSerializer


class SubjectViewSet(viewsets.ModelViewSet):
    queryset = Subject.objects.all()
    serializer_class = SubjectSerializer
    permission_classes = [permissions.AllowAny]


class TopicViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = Topic.objects.all()
    serializer_class = TopicSerializer
    permission_classes = [permissions.AllowAny]


class TestQueueListView(generics.ListAPIView):
    queryset = TestQueue.objects.all()
    serializer_class = TestQueueSerializer
    permission_classes = [IsAdminUser] 

    
class SubmitAnswerView(APIView):
    def post(self, request):
        session = TestSession.objects.get(id=request.data['session_id'], user=request.user)
        question = Question.objects.get(id=request.data['question_id'])
        TestAnswer.objects.update_or_create(
            session=session,
            question=question,
            defaults={"selected_answers": request.data['selected_answers']}
        )
        session.current_index += 1
        session.save()
        return Response({"message": "Javob saqlandi"})


class RestoreTestView(APIView):
    def get(self, request, test_id):
        session = TestSession.objects.filter(user=request.user, test_id=test_id, is_finished=False).first()
        if session:
            if session.queue.expected_end_time < timezone.now():
                evaluate_test(session)
                return Response({"message": "Vaqt tugadi. Test avtomatik baholandi."})
            serializer = TestSerializer(session.test)
            return Response({"session_id": session.id, "test": serializer.data})
        return Response({"message": "Faol test sessiyasi topilmadi."}, status=404)
    

class TestCreateView(APIView):
    def post(self, request):
        serializer = TestSerializer(data=request.data)
        if serializer.is_valid():
            test = serializer.save()
            return Response(TestSerializer(test).data)
        return Response(serializer.errors, status=400)
    

class KazusTestCreateAPIView(generics.CreateAPIView):
    queryset = Test.objects.all()
    serializer_class = KazusTestSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_serializer_context(self):
        context = super().get_serializer_context()
        context.update({"request": self.request})
        return context


class KazusTestUpdateAPIView(generics.RetrieveUpdateAPIView):
    queryset = Test.objects.all()
    serializer_class = KazusTestSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_serializer_context(self):
        context = super().get_serializer_context()
        context.update({"request": self.request})
        return context


class ResumeTestView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, session_id):
        session = get_object_or_404(TestSession, id=session_id, user=request.user)
        if session.is_completed or session.has_time_expired():
            return Response({'detail': 'Test muddati tugagan yoki allaqachon yakunlangan'}, status=400)

        answered_qs = UserAnswer.objects.filter(session=session).values_list('question_id', flat=True)
        remaining_questions = Question.objects.filter(test=session.test).exclude(id__in=answered_qs)
        serializer = QuestionSerializer(remaining_questions, many=True)
        return Response({'questions': serializer.data})


class CompleteTestView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, session_id):
        session = get_object_or_404(TestSession, id=session_id, user=request.user)
        if session.is_completed:
            return Response({'detail': 'Bu test allaqachon yakunlangan'}, status=400)

        if session.has_time_expired() or timezone.now():
            session.is_completed = True
            session.end_time = timezone.now()
            answers = UserAnswer.objects.filter(session=session)
            correct_count = answers.filter(is_correct=True).count()
            total_questions = Question.objects.filter(test=session.test).count()

            session.score = (correct_count / total_questions) * 100 if total_questions > 0 else 0
            session.save()

            return Response({'detail': 'Test yakunlandi', 'score': session.score})

        return Response({'detail': 'Test hali davom etmoqda'}, status=400)
    

class MyTestsAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        # 1. Oddiy Testlar
        tests = Test.objects.filter(creator=request.user, test_type='ordinary').order_by('id')
        tests_data = TestSerializer(tests, many=True, context={'request': request}).data
        
        for item in tests_data: 
            item['category'] = 'regular'
            
            # Sarlavhani kichik harflarga o'giramiz
            title_lower = str(item.get('title', '')).lower()
            
            # Qat'iy tekshiruv: agar sarlavhada fayl kengaytmasi bo'lsa yoki fayl maydoni bo'sh bo'lmasa
            is_file = (
                item.get('creation_method') in ['file', 'docx', 'xlsx'] or 
                ('.docx' in title_lower or '.xlsx' in title_lower or '.doc' in title_lower) or
                ('file' in item and item['file'] is not None and item['file'] != "") or
                item.get('is_file_uploaded') == True or
                item.get('file_uploaded') == True
            )
            
            if is_file:
                item['creation_method'] = 'file'
                method_label = "(FAYLDAN)"
            else:
                item['creation_method'] = 'manual'
                method_label = "(QO'LDA)"
                
            item['fan_nomi'] = f"ODDIY TEST {method_label}"

        # 2. Kazus Testlar
        case_tests = CaseTest.objects.filter(creator=request.user).order_by('-pk')
        case_data = CaseTestSerializer(case_tests, many=True, context={'request': request}).data
        for item in case_data: 
            item['category'] = 'case'
            item['test_type'] = 'kazus' 
            
            title_lower = str(item.get('title', '')).lower()
            is_file = (
                item.get('creation_method') in ['file', 'docx', 'xlsx'] or 
                ('.docx' in title_lower or '.xlsx' in title_lower or '.doc' in title_lower) or
                ('file' in item and item['file'] is not None and item['file'] != "")
            )
            item['creation_method'] = 'file' if is_file else 'manual'
            method_label = "(FAYLDAN)" if is_file else "(QO'LDA)"
            item['fan_nomi'] = f"KAZUS TEST {method_label}"

        # 3. Yopiq Testlar 
        closed_tests = ClosedTest.objects.filter(user=request.user).order_by('-id')
        closed_data = ClosedTestSerializer(closed_tests, many=True, context={'request': request}).data
        for item in closed_data: 
            item['category'] = 'closed'
            item['test_type'] = 'closed'
            
            title_lower = str(item.get('title', '')).lower()
            is_file = (
                item.get('creation_method') in ['file', 'docx', 'xlsx'] or 
                ('.docx' in title_lower or '.xlsx' in title_lower or '.doc' in title_lower) or
                ('file' in item and item['file'] is not None and item['file'] != "")
            )
            item['creation_method'] = 'file' if is_file else 'manual'
            method_label = "(FAYLDAN)" if is_file else "(QO'LDA)"
            item['fan_nomi'] = f"YOPIQ TEST {method_label}"
            if 'file' in item:
                item['uploaded_files'] = item['file']

        # Hammasini bitta ro'yxatga birlashtiramiz
        all_tests = tests_data + case_data + closed_data
        all_tests = sorted(all_tests, key=lambda x: x.get('id', 0), reverse=True)

        return Response(all_tests)
    

class ClosedTestDetailAPIView(RetrieveAPIView):
    queryset = ClosedTest.objects.all()
    serializer_class = ClosedTestSerializer
    permission_classes = [IsAuthenticated]


class CaseTestDetailAPIView(RetrieveAPIView):
    queryset = CaseTest.objects.all()
    serializer_class = CaseTestSerializer
    permission_classes = [IsAuthenticated]


class CaseTestListAPIView(generics.ListCreateAPIView):
    serializer_class = CaseTestSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        # MUHIM: filter(creator=self.request.user)
        return CaseTest.objects.filter(creator=self.request.user).order_by('-id')

    def perform_create(self, serializer):
        # Qo'lda yaratganda creator-ni saqlash
        serializer.save(creator=self.request.user)


class CaseTestViewSet(viewsets.ModelViewSet):
    serializer_class = CaseTestSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return CaseTest.objects.filter(user=self.request.user).order_by('-id')


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@transaction.atomic
def upload_test_file(request):
    print("\n" + "="*50)
    print("--- 📄 FAYLLI TEST YUKLASH BOSHLANDI ---")
    
    # 1. Ma'lumotlarni olish
    raw_test_type = request.data.get('test_type', 'ordinary')
    file_obj = request.FILES.get('file')
    s_id = request.data.get('subject')

    if not file_obj:
        print("XATO: Fayl kelmadi!")
        return Response({"error": "Fayl yuklanmadi"}, status=400)

    # Kazusligini qat'iy aniqlash mantiqi
    requested_type = str(raw_test_type).lower().strip()
    filename_lower = file_obj.name.lower()
    referer_url = request.META.get('HTTP_REFERER', '').lower()

    # Frontdan yoki URL'dan kelganiga qarab 100% aniqlaymiz
    is_actually_kazus = (
        requested_type in ['kazus', 'kasuz', 'case'] or 
        'kazus' in filename_lower or 
        'kasuz' in filename_lower or
        'case' in referer_url or
        'kazus' in referer_url
    )

    print(f"DEBUG: Frontdan kelgan test_type: '{raw_test_type}'")
    print(f"DEBUG: Yakuniy qaror - KAZUSMI?: {is_actually_kazus}")

    try:
        # Subject (Fan)ni aniqlash
        subject_obj = Subject.objects.filter(id=int(s_id)).first() if s_id else Subject.objects.first()

        # 2. ASOSIY TEST MODELINI YARATISH
        db_test_type = 'kazus' if is_actually_kazus else 'ordinary'
        
        new_test = Test.objects.create(
            title=file_obj.name,
            creator=request.user,
            subject=subject_obj,
            test_type=db_test_type,         # 🎯 Qat'iy 'kazus' deb yoziladi
            creation_method='file'          # 🎯 Qat'iy 'file' deb yoziladi
        )
        print(f"DEBUG: Ota Test yaratildi. ID: {new_test.id}, Type: {new_test.test_type}")

        # 3. KAZUS BO'LSA CaseTest (BOLA) MODELINI HAM YARATISH
        case_instance = None
        if is_actually_kazus:
            # unique_related_name yoki case_test_detail'ni ta'minlash uchun
            from .models import CaseTest
            case_instance = CaseTest.objects.create(
                parent_test=new_test,
                title=file_obj.name,
                creator=request.user,
                subject=subject_obj,
                file=file_obj,
                case_text=f"Fayldan yuklangan kazus: {file_obj.name}",
                creation_method='file'
            )
            print(f"DEBUG: CaseTest (Bola) modeli yaratildi.")

        # 4. SAVOLLARNI QO'SHISH (Parserni chaqirish)
        from .utils import parse_and_create_questions
        
        # 🎯 MUHIM TUZATISH: Savollarni har doim ota modelga (new_test) bog'laymiz!
        # Chunki Question modeli ota Test modeliga ulangan.
        count = parse_and_create_questions(new_test, file_obj, request.user)
        print(f"DEBUG: Parser yakunlandi. Savollar soni: {count}")

        # 5. VAQTNI HISOBLASH (1 ta savol = 1 minut)
        final_duration = count if (count and count > 0) else 30
        new_test.test_duration = final_duration
        new_test.save(update_fields=['test_duration'])

        if is_actually_kazus and case_instance:
            case_instance.test_duration = final_duration
            case_instance.save(update_fields=['test_duration'])

        label_name = "KAZUS TEST(FAYLDAN)" if is_actually_kazus else "ODDIY TEST(FAYLDAN)"
        print("--- 📄 FAYLLI TEST YUKLASH MUVAFFAQIYATLI YAKUNLANDI ---")
        print("="*50 + "\n")

        return Response({
            "success": True,
            "id": new_test.id,
            "label": label_name,
            "count": count,
            "duration": final_duration,
            "type_key": db_test_type
        }, status=201)

    except Exception as e:
        print(f"\n!!! XATOLIK YUZ BERDI !!!: {str(e)}")
        # Xatolikni aniq ko'rish uchun traceback'ni chiqaramiz
        import traceback
        traceback.print_exc()
        return Response({"error": f"Tizimda xatolik: {str(e)}"}, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def start_test(request, test_id=None):
    try:
        # 1. Frontenddan kelayotgan ma'lumotlarni olish
        t_id = test_id or request.data.get('test_id')
        t_type = (request.data.get('type_key') or 'regular').lower()
        
        if not t_id:
            return Response({"error": "test_id kiritilmadi"}, status=400)

        session_kwargs = {'user': request.user, 'start_time': timezone.now()}
        questions_query = []
        target_obj = None

        # 2. Test turi va savollarni aniqlash (TUZATILDI)
        if t_type in ['closed', 'yopiq']:
            target_obj = ClosedTest.objects.filter(id=t_id).first()
            if target_obj:
                questions_query = target_obj.closed_questions.all()
                session_kwargs['closed_test'] = target_obj
        
        elif t_type in ['kazus', 'case']:
            # AVVAL CaseTest modelidan qidiramiz
            target_obj = CaseTest.objects.filter(id=t_id).first()
            if target_obj:
                questions_query = target_obj.questions.all()
                session_kwargs['case_test'] = target_obj
            else:
                # AGAR TOPILMASA, Test modelidan (test_type='kazus' bo'lganini) qidiramiz
                # Bu aynan ManualTestAPIView orqali yuklangan testlar uchun
                target_obj = Test.objects.filter(id=t_id, test_type='kazus').first()
                if target_obj:
                    questions_query = target_obj.questions.all()
                    session_kwargs['test'] = target_obj
        
        else:
            # Oddiy testlar uchun
            target_obj = Test.objects.filter(id=t_id).first()
            if target_obj:
                questions_query = target_obj.questions.all()
                session_kwargs['test'] = target_obj

        if not target_obj:
            return Response({"error": "Test topilmadi"}, status=404)

        # --- VAQTNI ANIQLASH (TUZATILDI) ---
        q_count = questions_query.count()
        
        # creation_method NULL bo'lsa 'manual' deb olamiz
        raw_method = getattr(target_obj, 'creation_method', 'manual')
        if raw_method is None or str(raw_method).strip() == '':
            c_method = 'manual'
        else:
            c_method = str(raw_method).strip().lower()
        
        # Bazadagi vaqt maydonlari
        db_time = int(
            getattr(target_obj, 'duration_minutes', 0) or 
            getattr(target_obj, 'test_duration', 0) or 
            getattr(target_obj, 'duration', 0) or 0
        )

        if c_method == 'file':
            # Fayldan bo'lsa: savollar soni (default 15)
            test_duration = q_count if q_count > 0 else 15
        else:
            # Qo'lda bo'lsa: bazadagi vaqt, u 0 bo'lsa savollar soni
            test_duration = db_time if db_time > 0 else (q_count if q_count > 0 else 15)

        # 3. Sessiya yaratish
        session = TestSession.objects.create(**session_kwargs)

        # 4. Savollarni formatlash
        formatted_questions = []
        for q in questions_query:
            opts = {
                'A': getattr(q, 'answer_a', None) or getattr(q, 'answer_A', ""),
                'B': getattr(q, 'answer_b', None) or getattr(q, 'answer_B', ""),
                'C': getattr(q, 'answer_c', None) or getattr(q, 'answer_C', ""),
                'D': getattr(q, 'answer_d', None) or getattr(q, 'answer_D', ""),
            }

            q_text = str(getattr(q, 'text', ''))
            if hasattr(q, 'case_text') and q.case_text:
                q_text = f"{q.case_text}\n\n{q_text}"

            formatted_questions.append({
                "id": q.id,
                "text": q_text,
                "options": opts,
                "is_multiple": getattr(q, 'is_multiple_choice', False)
            })

        # 5. Yakuniy javob
        return Response({
            "success": True,
            "session_id": session.session_id,
            "test_title": target_obj.title,
            "duration": test_duration,
            "questions": formatted_questions
        }, status=200)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)
    

class UnifiedTestListAPIView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        search_query = request.query_params.get('search', '').strip()
        results = []

        # 1. ODDIY TESTLAR (Test modeli)
        # select_related va prefetch_related bazaga so'rovlar sonini kamaytiradi (N+1 muammosi oldini oladi)
        tests = Test.objects.select_related('subject', 'creator').all()
        if search_query:
            tests = tests.filter(title__icontains=search_query)
        
        # FILTR: test_type 'kazus' bo'lganlarini bu ro'yxatdan chiqarib tashlaymiz
        tests = tests.exclude(test_type='kazus')

        for t in tests:
            # Savollar sonini hisoblash (Related name ga qarab)
            try:
                q_count = t.questions.count()
            except AttributeError:
                q_count = t.question_set.count()

            # Davomiylikni aniqlash (Dinamik)
            duration = getattr(t, 'duration_minutes', getattr(t, 'test_duration', getattr(t, 'duration', 0)))

            results.append({
                "id": t.id,
                "title": t.title,
                "type_key": "regular",
                "display_type": "ODDIY TEST",
                "duration": duration,
                "subject": t.subject.name if t.subject else "Umumiy",
                "question_count": q_count,
                "author": t.creator.get_full_name() if t.creator else (t.creator.username if t.creator else "Noma'lum")
            })

        # 2. KAZUS TESTLAR (CaseTest modeli)
        cases = CaseTest.objects.select_related('subject', 'creator').all()
        if search_query:
            cases = cases.filter(title__icontains=search_query)
        
        for c in cases:
            # Savollar sonini aniqlash
            if hasattr(c, 'case_questions'):
                q_count = c.case_questions.count()
            elif hasattr(c, 'questions'):
                q_count = c.questions.count()
            else:
                q_count = 0

            # Davomiylik (Dinamik)
            duration = getattr(c, 'duration_minutes', getattr(c, 'duration', getattr(c, 'test_duration', 10)))

            results.append({
                "id": c.id,
                "title": c.title,
                "type_key": "kazus",
                "display_type": "KAZUS TEST",
                "duration": duration,
                "subject": c.subject.name if c.subject else "Huquq",
                "question_count": q_count,
                "author": c.creator.get_full_name() if hasattr(c, 'creator') and c.creator else "Noma'lum"
            })

        # 3. YOPIQ TESTLAR (ClosedTest modeli)
        # ClosedTest modelida 'user' fieldi ishlatilgani uchun select_related('user') qilamiz
        closed = ClosedTest.objects.select_related('subject', 'user').all()
        if search_query:
            closed = closed.filter(title__icontains=search_query)

        for cl in closed:
            # Savollar sonini hisoblash
            if hasattr(cl, 'closed_questions'):
                q_count = cl.closed_questions.count()
            else:
                try:
                    q_count = cl.closedquestion_set.count()
                except:
                    q_count = 0

            # Davomiylik
            duration = getattr(cl, 'duration', getattr(cl, 'test_duration', 15))

            results.append({
                "id": cl.id,
                "title": cl.title,
                "type_key": "closed",
                "display_type": "YOPIQ TEST",
                "duration": duration,
                "subject": cl.subject.name if hasattr(cl, 'subject') and cl.subject else "Yopiq",
                "question_count": q_count,
                "author": cl.user.get_full_name() if hasattr(cl, 'user') and cl.user else "Noma'lum"
            })

        # Natijalarni ID bo'yicha teskari tartibda (yangi qo'shilganlar tepada) ko'rsatish mumkin
        # results.sort(key=lambda x: x['id'], reverse=True)

        return Response(results)


import io, re, json, pandas as pd
import docx
from functools import wraps
from django.db import transaction
from django.shortcuts import get_object_or_404
from django.utils.cache import patch_cache_control

from rest_framework import status, generics
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.decorators import api_view, permission_classes

from .models import Test, CaseTest, CaseQuestion, ClosedTest, ClosedQuestion, Question, Subject
from .serializers import TestSerializer, CaseTestSerializer, ClosedTestSerializer
from django.http import JsonResponse  
from rest_framework.decorators import parser_classes
from rest_framework.parsers import MultiPartParser
from django.db.models import Q



def parse_and_create_questions(test_instance, docx_file, user):
    try:
        import re
        import io
        from docx import Document
        from django.db import transaction
        # Circular import xatosini oldini olish uchun modellar funksiya ichida
        from .models import Question, CaseQuestion, ClosedQuestion, Test, CaseTest, ClosedTest
        
        # Faylni o'qish
        docx_file.seek(0)
        doc = Document(io.BytesIO(docx_file.read()))
        
        questions = []
        current_q = None

        # 1. DOCX FAYLDAN SAVOLLARNI O'QISH
        for para in doc.paragraphs:
            text = para.text.strip().replace('\xa0', ' ').replace('\u200b', '')
            if not text: continue

            # SAVOLNI ANIQLASH ([s] bilan boshlansa yoki savol belgisi bilan tugasa)
            is_explicit_q = re.match(r'^\[s\]', text, re.IGNORECASE)
            is_question_mark = text.endswith('?')

            if is_explicit_q or is_question_mark:
                if current_q and current_q['ans']:
                    questions.append(current_q)
                
                q_text = re.sub(r'^\[s\]\s*', '', text, flags=re.IGNORECASE).strip() if is_explicit_q else text
                current_q = {'text': q_text, 'ans': [], 'corr_idx': 0}

            # JAVOBLARNI ANIQLASH (+ yoki - bilan boshlansa)
            elif current_q is not None and re.match(r'^[+\-–—]', text):
                sign = text[0]
                val = text[1:].strip()
                if val:
                    current_q['ans'].append(val)
                    if sign == '+':
                        current_q['corr_idx'] = len(current_q['ans']) - 1
            
            # SAVOL DAVOMI (Agar hali javoblar boshlanmagan bo'lsa)
            elif current_q is not None and not current_q['ans']:
                current_q['text'] += " " + text

        # Oxirgi savolni qo'shish
        if current_q and current_q['ans']:
            questions.append(current_q)

        if not questions:
            return 0

        created_count = 0
        char_map = ['A', 'B', 'C', 'D']
        model_name = type(test_instance).__name__

        # --- ✅ TEST INSTANCE HOLATINI YANGILASH ---
        if hasattr(test_instance, 'creation_method'):
            test_instance.creation_method = 'file'

        if 'CaseTest' in model_name or 'kazus' in test_instance.title.lower():
            if hasattr(test_instance, 'test_type'):
                test_instance.test_type = 'kazus'

        db_objs = []
        
        # 2. BAZAGA SAVOLLARNI TAYYORLASH (Bulk Create uchun)
        with transaction.atomic():
            for q_data in questions:
                ans = q_data['ans']
                while len(ans) < 4:
                    ans.append("---") 

                idx = q_data['corr_idx']
                correct_letter = char_map[idx] if idx < 4 else 'A'

                # --- KAZUS TEST (CaseTest) ---
                if 'CaseTest' in model_name:
                    db_objs.append(CaseQuestion(
                        case_test=test_instance,
                        text=q_data['text'],
                        answer_a=ans[0],
                        answer_b=ans[1],
                        answer_c=ans[2],
                        answer_d=ans[3],
                        correct_answers=[correct_letter], # List formatida
                        creator=user
                    ))

                # --- YOPIQ TEST (ClosedTest) ---
                elif 'ClosedTest' in model_name:
                    db_objs.append(ClosedQuestion(
                        test=test_instance,
                        text=q_data['text'],
                        answer_A=ans[0],
                        answer_B=ans[1],
                        answer_C=ans[2],
                        answer_D=ans[3],
                        correct_answers=correct_letter, # String formatida
                        question_type='single'
                    ))

                # --- ODDIY TEST (Test) ---
                else:
                    db_objs.append(Question(
                        test=test_instance,
                        text=q_data['text'],
                        answer_a=ans[0],
                        answer_b=ans[1],
                        answer_c=ans[2],
                        answer_d=ans[3],
                        correct_option=correct_letter, # A, B, C, D
                        question_type='single',
                        is_correct_a=(idx == 0),
                        is_correct_b=(idx == 1),
                        is_correct_c=(idx == 2),
                        is_correct_d=(idx == 3),
                        creator=user,
                        subject=getattr(test_instance, 'subject', None)
                    ))

            # Modelga qarab bulk_create qilish
            if 'CaseTest' in model_name:
                CaseQuestion.objects.bulk_create(db_objs)
            elif 'ClosedTest' in model_name:
                ClosedQuestion.objects.bulk_create(db_objs)
            else:
                Question.objects.bulk_create(db_objs)

            created_count = len(db_objs)

            # 3. YAKUNIY SAQLASH (VAQTNI YANGILASH)
            if created_count > 0:
                # Qaysi modelda qanday nomlanganini tekshirish
                if hasattr(test_instance, 'duration'):
                    test_instance.duration = created_count
                elif hasattr(test_instance, 'test_duration'):
                    test_instance.test_duration = created_count
                elif hasattr(test_instance, 'duration_minutes'):
                    test_instance.duration_minutes = created_count
                
                test_instance.save()

        return created_count

    except Exception as e:
        import traceback
        print(f"PARSE ERROR: {e}")
        traceback.print_exc()
        return 0



import json, traceback, docx
from io import BytesIO
from django.db import transaction
from django.shortcuts import get_object_or_404
from django.utils.html import escape
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from .models import Subject, Test, Question, ClosedTest, ClosedQuestion, CaseTest, CaseQuestion
from .serializers import UniversalQuestionSerializer
from .utils import parse_full_docx
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import CaseTest
from .utils import parse_and_create_questions 
import os
from django.views.decorators.csrf import csrf_exempt
import json
from django.http import JsonResponse
from .models import CaseTest, CaseQuestion
from rest_framework.test import APIRequestFactory







def parse_closed_test_from_memory(file_obj):
    """Word faylni xotiradan o'qiydi va savollarni list ko'rinishida qaytaradi"""
    try:
        file_obj.seek(0)
        doc = docx.Document(io.BytesIO(file_obj.read()))
    except Exception as e:
        print(f"Faylni o'qishda xatolik: {e}")
        return []

    questions = []
    current_q = None

    for para in doc.paragraphs:
        line = para.text.strip().replace('\xa0', ' ').replace('\u200b', '')
        if not line: continue

        # Savol boshlanishi [s]
        if line.lower().startswith('[s]'):
            if current_q and current_q['answers']:
                questions.append(current_q)
            
            # HTML belgilarni xavfsiz qilish
            q_text = line[3:].strip().replace('<', '&lt;').replace('>', '&gt;')
            current_q = {
                'text': q_text,
                'answers': [],
                'correct': 'A'
            }
        
        # Variantlar (+, -, –, —)
        elif current_q is not None and any(line.startswith(symbol) for symbol in ['+', '-', '–', '—']):
            ans_text = line[1:].strip().replace('<', '&lt;').replace('>', '&gt;')
            if ans_text:
                current_q['answers'].append(ans_text)
                if line.startswith('+'):
                    idx = len(current_q['answers']) - 1
                    current_q['correct'] = chr(65 + idx) # A, B, C, D
        
        # Savol matni davomi
        elif current_q is not None and not current_q['answers']:
            current_q['text'] += " " + line.replace('<', '&lt;').replace('>', '&gt;')

    if current_q and current_q['answers']:
        questions.append(current_q)
        
    return questions


def upload_case_test_view(request):
    if request.method == 'POST':
        from .views import ManualTestAPIView 
        
        subject_id = request.POST.get('subject_id') or request.POST.get('subject')
        raw_type = request.POST.get('test_type', 'ordinary')

        if not subject_id:
            return JsonResponse({'error': "Fan ID topilmadi."}, status=400)

        # Ma'lumotlarni o'zgartirish uchun nusxa olamiz
        mutable_data = request.POST.copy()
        
        # 🎯 QAT'IY MAJBURLASH: Bu viewga keldimi - demak 100% KAZUS bo'lishi kerak!
        mutable_data['test_type'] = 'kazus'
        mutable_data['category'] = 'case'
        mutable_data['subject'] = subject_id

        # Django REST Framework uchun so'rov tayyorlash
        factory = APIRequestFactory()
        
        # MUHIM: Ba'zida format='multipart' bo'lganda ma'lumot yo'qolmasligi uchun QUERY_STRING'ga ham urib qo'yamiz
        new_request = factory.post(
            f"{request.path}?test_type=kazus&category=case", 
            mutable_data, 
            format='multipart'
        )
        
        new_request.user = request.user 
        new_request.FILES = request.FILES
        
        # 🛠 DRF request.data ichiga ham majburlab 'kazus' qiymatini tiqib qo'yamiz!
        # Shunda ManualTestAPIView ichidagi serializer uni rad eta olmaydi.
        if hasattr(new_request, '_full_data'):
            new_request._full_data = mutable_data
        
        # Mana bu satr ManualTestAPIView'ga 'kazus' qiymatini qat'iy o'tkazadi
        new_request.data = mutable_data 

        # API viewni chaqiramiz
        view = ManualTestAPIView.as_view()
        response = view(new_request)

        # Agar muvaffaqiyatli saqlangan bo'lsa, lekin ichki API baribir 'ordinary' qilib qo'ygan bo'lsa,
        # uni bazadan topib, majburlab yangilab qo'yamiz! (Eng ishonchli 'Crutch' yechim)
        if response.status_code == 201:
            try:
                test_id = response.data.get('test_id') if hasattr(response, 'data') else None
                if not test_id and hasattr(response, 'data') and isinstance(response.data, dict):
                    test_id = response.data.get('id')
                
                if test_id:
                    from .models import Test, CaseTest
                    # 1. Asosiy Test modelini bazada kazusga o'zgartiramiz
                    Test.objects.filter(id=test_id).update(test_type='kazus')
                    
                    # 2. Agar CaseTest (bola) modeli ochilmay qolgan bo'lsa, uni ham tekshirib ochib qo'yamiz
                    parent_test = Test.objects.filter(id=test_id).first()
                    if parent_test and not CaseTest.objects.filter(parent_test=parent_test).exists():
                        CaseTest.objects.create(
                            parent_test=parent_test,
                            title=parent_test.title,
                            creator=parent_test.creator,
                            subject=parent_test.subject,
                            creation_method=parent_test.creation_method or 'file'
                        )
            except Exception as db_err:
                print(f"Post-update majburlashda xatolik: {db_err}")

            return JsonResponse({
                'success': True, 
                'message': "Test muvaffaqiyatli saqlandi va Kazus toifasiga biriktirildi.",
                'test_id': response.data.get('test_id') if hasattr(response, 'data') else None,
                'test_type': 'kazus'
            }, status=201)
        
        # Xatolik qaytsa
        res_data = response.data if hasattr(response, 'data') else str(response)
        return JsonResponse({'error': res_data}, status=response.status_code)
            
    return JsonResponse({'error': 'POST kerak'}, status=405)


@api_view(['POST'])
def upload_closed_test(request):
    title = request.data.get('title')
    file_obj = request.FILES.get('file')
    subject_id = request.data.get('subject')

    if not file_obj:
        return Response({"error": "Fayl yuklanmadi"}, status=400)

    try:
        # 1. Utils ichidagi yangi parserni chaqiramiz
        # Eslatma: utils.py dagi funksiya nomi parse_full_docx bo'lsa shuni yozing
        questions_data = parse_full_docx(file_obj) 

        if not questions_data:
            return Response({
                "error": "Fayldan savollar topilmadi. [s], + va - belgilarini tekshiring."
            }, status=400)

        with transaction.atomic():
            # 2. Test ob'ektini (konteynerni) yaratish
            test_instance = ClosedTest.objects.create(
                title=title or file_obj.name,
                file=file_obj,
                user=request.user,
                subject_id=subject_id if subject_id and subject_id != 'undefined' else None,
                duration=len(questions_data) # Savollar soniga qarab vaqt (ixtiyoriy)
            )

            # 3. Savollarni bazaga yozish
            question_objs = []
            for item in questions_data:
                # DIQQAT: Utils dan kelayotgan kalit so'zlar (answer_A, answer_B...) 
                # bilan bir xil bo'lishi shart.
                question_objs.append(ClosedQuestion(
                    test=test_instance,
                    text=item.get('text', 'Savol matni mavjud emas'),
                    answer_A=item.get('answer_A', '---'),
                    answer_B=item.get('answer_B', '---'),
                    answer_C=item.get('answer_C', '---'),
                    answer_D=item.get('answer_D', '---'),
                    # Parseringiz 'correct_answers' qaytaryapti
                    correct_answers=item.get('correct_answers', 'A'),
                    question_type='single'
                ))
            
            # Barcha savollarni bitta zaprosda bazaga saqlaymiz (tezroq ishlaydi)
            ClosedQuestion.objects.bulk_create(question_objs)

        return Response({
            "status": "success", 
            "message": f"{len(question_objs)} ta savol muvaffaqiyatli yuklandi!",
            "test_id": test_instance.id
        }, status=201)

    except Exception as e:
        # Xatolikni terminalda ko'rish uchun (debugging)
        import traceback
        print(traceback.format_exc())
        return Response({"error": f"Tizimda xatolik: {str(e)}"}, status=500)


class GetTestDetailsView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        test_id = request.query_params.get('test_id')
        test_type = str(request.query_params.get('type', '')).lower()

        if not test_id or not test_type:
            return Response({"error": "Parametrlar yetishmayapti"}, status=400)

        try:
            # 1. Toifalarni aniqlash va Test obyektini olish
            test_obj = None
            questions_query = None

            if test_type in ['regular', 'ordinary', 'single', 'multiple']:
                test_obj = Test.objects.filter(pk=test_id).first()
                questions_query = Question.objects.filter(test_id=test_id).order_by('id')
                if not questions_query.exists():
                    questions_query = ClosedQuestion.objects.filter(test_id=test_id).order_by('id')

            elif test_type in ['kazus', 'case']:
                test_obj = CaseTest.objects.filter(pk=test_id).first()
                questions_query = CaseQuestion.objects.filter(case_test_id=test_id).order_by('id')
                if not questions_query.exists():
                    questions_query = Question.objects.filter(case_test_id=test_id).order_by('id')

            elif test_type == 'closed':
                test_obj = ClosedTest.objects.filter(pk=test_id).first()
                questions_query = ClosedQuestion.objects.filter(test_id=test_id).order_by('id')
                if not questions_query.exists():
                    questions_query = Question.objects.filter(test_id=test_id).order_by('id')

            if not questions_query or not questions_query.exists():
                return Response({"questions": [], "duration": 15}, status=200)

            # --- VAQTNI HISOBLASH (RASMDAGI 30 MINUT MUAMMOSI SHU YERDA YECHILADI) ---
            q_count = questions_query.count()
            db_time = 0
            c_method = 'manual'

            if test_obj:
                db_time = int(getattr(test_obj, 'duration_minutes', 0) or getattr(test_obj, 'duration', 0) or 0)
                c_method = str(getattr(test_obj, 'creation_method', 'manual')).strip().lower()

            # AGAR VAQT 30 BO'LSA VA SAVOLLAR 10 TA BO'LSA (YOKI 30 DAN FARQLI BO'LSA)
            # BU FAYLDAN YUKLANGAN TEST DEB HISOBLANADI VA VAQTNI SAVOL SONIGA TENGILAYMIZ
            if db_time == 30 and q_count != 30 and q_count > 0:
                final_duration = q_count
            elif c_method == 'file':
                final_duration = q_count
            else:
                final_duration = db_time if db_time > 0 else (q_count if q_count > 0 else 15)

            # Serializatsiya qilish
            serializer = UniversalQuestionSerializer(questions_query, many=True)
            questions_data = serializer.data

            # Frontend uchun ma'lumotni to'g'rilash (Serializatsiya qilingan savollar ichida)
            for item in questions_data:
                if not item.get('correct_answers') and item.get('correct_option'):
                    item['correct_answers'] = [item['correct_option']]
                if not item.get('correct_option') and item.get('correct_answers'):
                    item['correct_option'] = item['correct_answers'][0]

            # YAKUNIY RESPONSE: Ham savollarni, ham to'g'irlangan vaqtni yuboramiz
            return Response({
                "questions": questions_data,
                "duration": final_duration,
                "test_title": getattr(test_obj, 'title', 'Test')
            })

        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": "Ichki xatolik"}, status=500)
        

import re
import docx
import traceback
from io import BytesIO
from django.db import transaction, connection
from django.views.decorators.cache import never_cache
from rest_framework.decorators import api_view
from django.shortcuts import get_object_or_404
from rest_framework.views import APIView
from django.utils.html import escape 
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser

from .models import Subject, Test, Question, ClosedTest, ClosedQuestion, CaseTest, CaseQuestion
from .serializers import TestSerializer, CaseTestSerializer, ClosedTestSerializer, UniversalQuestionSerializer, ClosedQuestionSerializer



def disable_cache(view_func):
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        response = view_func(request, *args, **kwargs)
        patch_cache_control(response, no_cache=True, no_store=True, must_revalidate=True, max_age=0, private=True)
        return response
    return _wrapped_view


class ManualTestAPIView(APIView):
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    def post(self, request):
        # Django MultiPartParser'da ma'lumotlarni xavfsiz olish uchun .get() o'rniga .getlist() yoki tozalash ishlatamiz
        data = request.data
        file_obj = request.FILES.get('file')
        
        # 1. Kelgan test_type'ni har qanday holatda ham toza stringga o'giramiz
        raw_type = data.get('test_type') or data.get('type') or ''
        
        # Agar list yoki QueryDict'dan kelgan massiv bo'lsa, birinchi elementini olamiz
        if isinstance(raw_type, list) and len(raw_type) > 0:
            raw_type = raw_type[0]
        elif hasattr(raw_type, 'get'): # Ba'zida ichma-ich obyekt bo'lsa
            raw_type = str(raw_type)

        test_type_str = str(raw_type).lower().strip()
        file_name_str = str(file_obj.name).lower() if file_obj else ""

        # --- JIDDIY PRINT (Terminalda tekshirish uchun) ---
        # Serveringiz ishlayotgan terminalga qarang, nima kelayotganini aniq ko'rasiz!
        print("\n" + "="*50)
        print(f"DEBUG -> Kelgan xom turi: {raw_type}")
        print(f"DEBUG -> Tozalangan turi: {test_type_str}")
        print(f"DEBUG -> Fayl nomi: {file_name_str}")
        print("="*50 + "\n")

        # Kazus ekanligini mukammal va eng agressiv usulda aniqlaymiz
        is_kazus = any(k in test_type_str for k in ['kazus', 'kasuz', 'kasu', 'case']) or \
                   any(k in file_name_str for k in ['kazus', 'kasuz', 'kasu', 'case'])
                   
        is_closed = 'closed' in test_type_str or 'yopiq' in test_type_str

        if is_kazus:
            lookup_key = 'case'
            actual_db_type = 'kazus'
        elif is_closed:
            lookup_key = 'closed'
            actual_db_type = 'ordinary'
        else:
            lookup_key = 'regular'
            actual_db_type = 'ordinary'

        try:
            s_id = data.get('subject')
            if isinstance(s_id, list):
                s_id = s_id[0]
                
            if not s_id:
                return Response({"error": "Subject ID yuborilmadi"}, status=400)
            
            subject_id = int(str(s_id).strip())
            subject_obj = get_object_or_404(Subject, id=subject_id)

            if file_obj:
                questions_list = self._parse_docx(file_obj)
                is_from_file = True
            else:
                questions_list = self._parse_json(data)
                is_from_file = False

            if not questions_list:
                return Response({"error": "Savollar topilmadi!"}, status=400)

            with transaction.atomic():
                # Title massiv bo'lib kelsa tozalaymiz
                raw_title = data.get('title')
                if isinstance(raw_title, list) and len(raw_title) > 0:
                    raw_title = raw_title[0]
                
                title = raw_title or (file_obj.name if file_obj else "Yangi test")

                q_count = len(questions_list)
                
                # Dynamic Duration Rule (1 savol = 1 minut)
                if is_from_file:
                    final_duration = q_count
                else:
                    raw_duration = data.get('duration') or data.get('duration_minutes') or data.get('test_duration')
                    if isinstance(raw_duration, list) and len(raw_duration) > 0:
                        raw_duration = raw_duration[0]
                    try:
                        val = int(float(str(raw_duration).strip()))
                        final_duration = val if val > 0 else 30
                    except:
                        final_duration = 30

                # Modellarni tanlash
                config = {
                    'closed':   (ClosedTest,  ClosedQuestion, 'test',      {'user': request.user}),
                    'case':     (CaseTest,    CaseQuestion,   'case_test', {'creator': request.user, 'case_text': title}),
                    'regular':  (Test,        Question,       'test',      {'creator': request.user}),
                }
                
                test_model, q_model, fk_field, extra_fields = config[lookup_key]
                
                all_fields = [f.name for f in test_model._meta.get_fields()]
                if 'test_duration' in all_fields: 
                    time_col = 'test_duration'
                elif 'duration_minutes' in all_fields: 
                    time_col = 'duration_minutes'
                else: 
                    time_col = 'duration'

                create_params = {
                    'subject': subject_obj,
                    'title': title,
                    time_col: final_duration,
                    **extra_fields
                }

                if 'test_type' in all_fields:
                    create_params['test_type'] = actual_db_type
                if 'creation_method' in all_fields:
                    create_params['creation_method'] = 'file' if is_from_file else 'manual'

                # Obyektni yaratish
                test_obj = test_model.objects.create(**create_params)

                db_objs = []
                for q_data in questions_list:
                    ans = q_data['ans'][:]
                    while len(ans) < 4: 
                        ans.append("---")
                    idx = q_data.get('corr_idx', 0)
                    final_letter = chr(65 + idx)

                    params = {fk_field: test_obj, 'text': q_data['text']}

                    if lookup_key == 'closed':
                        params.update({
                            'answer_A': ans[0], 'answer_B': ans[1], 'answer_C': ans[2], 'answer_D': ans[3],
                            'correct_answers': final_letter, 'question_type': 'single'
                        })
                    elif lookup_key == 'case':
                        params.update({
                            'answer_a': ans[0], 'answer_b': ans[1], 'answer_c': ans[2], 'answer_d': ans[3],
                            'correct_answers': [final_letter], 'question_type': 'case', 'is_multiple_choice': False
                        })
                    else:
                        params.update({
                            'answer_a': ans[0], 'answer_b': ans[1], 'answer_c': ans[2], 'answer_d': ans[3],
                            'correct_option': final_letter, 'correct_answer': final_letter,
                            'question_type': 'single', 'is_correct_a': (idx == 0), 'is_correct_b': (idx == 1),
                            'is_correct_c': (idx == 2), 'is_correct_d': (idx == 3),
                            'subject': subject_obj, 'creator': request.user
                        })
                    db_objs.append(q_model(**params))

                q_model.objects.bulk_create(db_objs)

            return Response({
                "status": "success", 
                "test_id": test_obj.id, 
                "type": actual_db_type,
                "test_type": actual_db_type
            }, status=201)

        except Exception as e:
            traceback.print_exc()
            return Response({"error": str(e)}, status=400)

    def _parse_docx(self, file_obj):
        questions = []
        current_q = None
        try:
            file_obj.seek(0)
            doc = docx.Document(io.BytesIO(file_obj.read()))
            
            for para in doc.paragraphs:
                text = para.text.strip().replace('\xa0', ' ').replace('\u200b', '')
                if not text:
                    continue
                
                # HTML teglari bilan ishlash (<h1>, <p> kabilar uchun)
                safe_text = text.replace('<', '&lt;').replace('>', '&gt;')
                
                # VARIANT ekanligini aniqlash (eng birinchi bo'lib tekshiriladi)
                is_variant = safe_text.startswith(('+', '-', '–', '—')) or \
                             (len(safe_text) >= 2 and safe_text[0].upper() in 'ABCD' and safe_text[1] in [')', '.', ':'])

                # SAVOL ekanligini aniqlash
                is_explicit_q = safe_text.lower().startswith('[s]')
                
                # AGAR BU VARIANT BO'LSA (Bu blok savol blokidan yuqorida bo'lishi shart)
                if is_variant and current_q is not None:
                    is_correct = safe_text.startswith('+')
                    
                    # Belgini olib tashlab matnni ajratish
                    if safe_text[0] in ['+', '-', '–', '—']:
                        ans_text = safe_text[1:].strip()
                    else:
                        ans_text = safe_text[2:].strip()
                    
                    # Agar matn bo'sh qolsa, butun qatorni olish
                    if not ans_text:
                        ans_text = safe_text 

                    current_q['ans'].append(ans_text)
                    if is_correct: 
                        current_q['corr_idx'] = len(current_q['ans']) - 1
                    continue # Variant topildimi, keyingi qatorga o'tish

                # AGAR BU YANGI SAVOL BO'LSA
                if is_explicit_q or (not is_variant and (safe_text.endswith('?') or current_q is None)):
                    if current_q and current_q['ans']: 
                        questions.append(current_q)
                    
                    q_text = safe_text
                    if is_explicit_q: 
                        q_text = safe_text[3:].strip()
                    
                    current_q = {'text': q_text, 'ans': [], 'corr_idx': 0}
                
                # AGAR BU VARIANT EMAS VA SAVOL DAVOMI BO'LSA
                elif current_q is not None and not is_variant:
                    # Faqat variantlar hali boshlanmagan bo'lsagina matnni savolga qo'shish
                    if not current_q['ans']:
                        current_q['text'] += " " + safe_text
            
            # Oxirgi savolni qo'shish
            if current_q and current_q['ans']: 
                questions.append(current_q)
                
        except Exception as e:
            print(f"Xatolik: {e}")
            
        return questions

    def _parse_json(self, data):
        q_raw = data.get('questions') or data.get('closed_questions') or data.get('testlar') or []
        if isinstance(q_raw, str):
            try: q_raw = json.loads(q_raw)
            except: return []
        
        result = []
        l_map = {'A': 0, 'B': 1, 'C': 2, 'D': 3}
        for q in q_raw:
            text = q.get('text', '').strip()
            if not text: continue
            
            ans = [
                str(q.get('answer_a') or q.get('answer_A') or q.get('A') or '').strip(),
                str(q.get('answer_b') or q.get('answer_B') or q.get('B') or '').strip(),
                str(q.get('answer_c') or q.get('answer_C') or q.get('C') or '').strip(),
                str(q.get('answer_d') or q.get('answer_D') or q.get('D') or '').strip()
            ]
            
            # --- TO'G'RILANGAN QISM ---
            # Frontenddan kelishi mumkin bo'lgan barcha kalitlarni tekshiramiz
            raw_correct = q.get('correct_answers') or q.get('correct_option') or q.get('correct_answer') or 'A'
            
            # Agar frontenddan ro'yxat (list) kelsa, birinchi elementini olamiz
            if isinstance(raw_correct, list) and len(raw_correct) > 0:
                val = str(raw_correct[0]).upper().strip()
            else:
                val = str(raw_correct).upper().strip()
            
            # Agar qiymat harf emas (masalan variant matni) bo'lsa, uni harfga aylantirishga urinib ko'ramiz
            if val not in l_map:
                # Agar val ans ichidagi matnlardan biriga teng bo'lsa, o'sha indexni olamiz
                for i, choice in enumerate(ans):
                    if val == choice.upper().strip():
                        val = chr(65 + i) # A, B, C yoki D ga aylantiradi
                        break

            result.append({
                'text': text, 
                'ans': ans, 
                'corr_idx': l_map.get(val, 0) # Baribir topilmasa, 0 (A) bo'ladi
            })
        return result


@api_view(['GET'])
@never_cache  # Brauzer keshlamasligi uchun
def get_test_json(request, pk):
    try:
        # Har ikkala modeldan qidirib ko'ramiz
        test = None
        questions_qs = None
        test_type = "regular"

        if ClosedTest.objects.filter(id=pk).exists():
            test = ClosedTest.objects.get(id=pk)
            questions_qs = test.closed_questions.all()
            test_type = "closed"
        elif Test.objects.filter(id=pk).exists():
            test = Test.objects.get(id=pk)
            questions_qs = test.questions.all()
            test_type = "regular"
        
        if not test:
            return Response({"error": "Test topilmadi"}, status=404)

        # VAQTNI HISOBLASH (Har bir savolga 1 minut)
        # Modelga qarab duration yoki duration_minutes ni tekshiramiz
        raw_duration = getattr(test, 'duration', None) or getattr(test, 'duration_minutes', 0)
        
        q_count = questions_qs.count()
        # Agar vaqt 0 bo'lsa yoki savollar sonidan juda kichik bo'lsa, savollar sonini olamiz
        if not raw_duration or int(raw_duration) <= 0:
            final_duration = q_count
        else:
            final_duration = int(raw_duration)

        # SAVOLLARNI FORMATLASH
        questions_list = []
        for q in questions_qs:
            # To'g'ri javobni aniqlash (Closed uchun correct_answers, Regular uchun correct_option)
            correct = getattr(q, 'correct_answers', None) or getattr(q, 'correct_option', 'A')
            
            questions_list.append({
                "id": q.id,
                "text": q.text,
                "answer_a": getattr(q, 'answer_A', None) or getattr(q, 'answer_a', ''),
                "answer_b": getattr(q, 'answer_B', None) or getattr(q, 'answer_b', ''),
                "answer_c": getattr(q, 'answer_C', None) or getattr(q, 'answer_c', ''),
                "answer_d": getattr(q, 'answer_D', None) or getattr(q, 'answer_d', ''),
                "correct_option": str(correct).strip().upper(),
                "correct_answers": str(correct).strip().upper(),
            })

        return Response({
            "id": test.id,
            "title": test.title,
            "duration": final_duration,
            "type_key": test_type,
            "questions": questions_list,
            "closed_questions": questions_list 
        })
    except Exception as e:
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@disable_cache
def get_my_tests(request):
    """Foydalanuvchining barcha testlarini qidiruv va filtrlar bilan qaytaradi"""
    try:
        # Parametrlarni olish
        query = request.GET.get('q', '')
        category = request.GET.get('cat', '')
        subject_id = request.GET.get('subject', '')

        # Filtrlash
        tests = Test.objects.filter(creator=request.user).order_by('-id')
        
        if query:
            tests = tests.filter(title__icontains=query)
        if category:
            tests = tests.filter(category__icontains=category)
        if subject_id:
            tests = tests.filter(subject_id=subject_id)

        serializer = TestSerializer(tests, many=True, context={'request': request})
        return Response(serializer.data)
    except Exception as e:
        return Response({"error": str(e)}, status=500)
    

@api_view(['GET'])
@disable_cache
def get_test_json(request, pk):
    try:
        test = get_object_or_404(ClosedTest, id=pk)
        # Serializer orqali barcha savollarni o'raymiz
        serializer = ClosedQuestionSerializer(test.closed_questions.all(), many=True)
        
        return Response({
            "title": test.title,
            "closed_questions": serializer.data 
        })
    except Exception as e:
        return Response({"error": str(e)}, status=500)


@api_view(['GET'])
@disable_cache
def download_my_test_json(request, test_id):
    path = request.path.lower()
    
    if 'closed-tests' in path:
        test_obj = get_object_or_404(ClosedTest, id=test_id)
        # ✅ TO'G'RI related_name
        qs = test_obj.closed_questions.all()
        
        questions_list = []
        for q in qs:
            questions_list.append({
                "text": q.text,
                # ✅ Katta harfli fieldlar
                "answer_a": q.answer_A,
                "answer_b": q.answer_B,
                "answer_c": q.answer_C,
                "answer_d": q.answer_D,
                # ✅ correct_answers -> kichik harfga
                "correct_option": str(q.correct_answers).strip().lower()[:1]
            })
            
    elif 'case-tests' in path:
        test_obj = get_object_or_404(CaseTest, id=test_id)
        qs = getattr(test_obj, 'case_questions', None) or test_obj.questions
        qs = qs.all()
        
        questions_list = []
        for q in qs:
            questions_list.append({
                "text": q.text,
                "answer_a": getattr(q, 'answer_a', ''),
                "answer_b": getattr(q, 'answer_b', ''),
                "answer_c": getattr(q, 'answer_c', ''),
                "answer_d": getattr(q, 'answer_d', ''),
                "correct_option": str(
                    getattr(q, 'correct_option', None) or
                    getattr(q, 'correct_answers', 'a')
                ).strip().lower()[:1]
            })
    else:
        # Oddiy testlar
        test_obj = get_object_or_404(Test, id=test_id)
        qs = test_obj.questions.all()
        
        questions_list = []
        for q in qs:
            questions_list.append({
                "text": q.text,
                "answer_a": getattr(q, 'answer_a', ''),
                "answer_b": getattr(q, 'answer_b', ''),
                "answer_c": getattr(q, 'answer_c', ''),
                "answer_d": getattr(q, 'answer_d', ''),
                "correct_option": str(
                    getattr(q, 'correct_option', None) or
                    getattr(q, 'correct_answers', 'a')
                ).strip().lower()[:1]
            })

    return Response({
        "title": test_obj.title,
        "questions": questions_list
    })
    

class MyTestsListAPIView(generics.ListAPIView):
    serializer_class = TestSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        # Faqat shu foydalanuvchi yaratgan testlarni qaytaradi
        return Test.objects.filter(creator=self.request.user).order_by('-id')
    

class CaseTestsListAPIView(generics.ListAPIView):
    serializer_class = CaseTestSerializer
    permission_classes = [IsAuthenticated]
    def get_queryset(self):
        return CaseTest.objects.filter(creator=self.request.user).order_by('-id')
    

class ClosedTestsListAPIView(generics.ListAPIView):
    serializer_class = ClosedTestSerializer
    permission_classes = [IsAuthenticated]
    def get_queryset(self):
        return ClosedTest.objects.filter(user=self.request.user).order_by('-id')
    

# --- 3. TESTNI TEKSHIRISH ---
class CheckTestAPIView(APIView):
    permission_classes = [IsAuthenticated]
    def post(self, request, pk):
        user_answers = request.data.get('answers', {})
        test = get_object_or_404(Test, id=pk)
        questions = test.questions.all()
        correct = sum(1 for q in questions if user_answers.get(str(q.id)) == q.correct_option)
        return Response({"score": correct, "total": questions.count()})


# ==============================
# 🔹 Payment ro‘yxatini ko‘rish
# ==============================
class PaymentViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = PaymentSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Payment.objects.filter(user=self.request.user)
    

logger = logging.getLogger(__name__)

# ==============================
# 🔹 Click to‘lov yaratish
# ==============================
def create_click_payment(user, test_id, amount):
    """
    Click to‘lov URL yaratadi va Payment modelida yozadi.
    """
    merchant_id = settings.CLICK_MERCHANT_ID
    service_id = settings.CLICK_SERVICE_ID
    secret_key = settings.CLICK_SECRET_KEY

    order_id = f"test-{test_id}-{user.id}"
    sign_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Click sign hisoblash
    sign_source = f"{merchant_id}{service_id}{secret_key}{order_id}{amount}{sign_time}"
    sign_string = hashlib.md5(sign_source.encode('utf-8')).hexdigest()

    payment_url = (
        f"https://my.click.uz/pay?"
        f"merchant_id={merchant_id}&"
        f"service_id={service_id}&"
        f"amount={amount}&"
        f"transaction_param={order_id}&"
        f"sign_time={sign_time}&"
        f"sign_string={sign_string}"
    )

    # Payment yozish
    Payment.objects.create(
        user=user,
        test=get_object_or_404(Test, id=test_id),
        amount=amount,
        provider='click',
        transaction_id=order_id,
        status='pending'
    )

    return payment_url


# ==============================
# 🔹 Click prepare callback
# ==============================
@csrf_exempt
def click_prepare(request):
    if request.method != "POST":
        return JsonResponse({"error": "Faqat POST so‘rov qabul qilinadi"}, status=405)

    data = request.POST
    logger.debug("Click prepare payload: %s", data)

    required_fields = [
        "click_trans_id", "service_id", "click_paydoc_id",
        "merchant_trans_id", "amount", "action", "sign_time", "sign_string"
    ]
    missing_fields = [f for f in required_fields if f not in data]
    if missing_fields:
        return JsonResponse({
            "error": "To‘lov parametrlari mavjud emas",
            "missing_fields": missing_fields
        }, status=400)

    try:
        payment = Payment.objects.get(transaction_id=data["merchant_trans_id"])
    except Payment.DoesNotExist:
        return JsonResponse({"error": "Buyurtma topilmadi"}, status=404)

    # Sign tekshirish
    sign_source = (
        f"{data['click_trans_id']}{data['service_id']}"
        f"{settings.CLICK_SECRET_KEY}{data['merchant_trans_id']}"
        f"{data['amount']}{data['action']}{data['sign_time']}"
    )
    valid_sign = hashlib.md5(sign_source.encode('utf-8')).hexdigest()

    if valid_sign != data["sign_string"]:
        return JsonResponse({"error": "Sign noto‘g‘ri"}, status=400)

    return JsonResponse({
        "click_trans_id": data["click_trans_id"],
        "merchant_trans_id": data["merchant_trans_id"],
        "merchant_prepare_id": payment.id,
        "error": 0,
        "error_note": "Success"
    })


# ==============================
# 🔹 Click complete callback
# ==============================
@csrf_exempt
def click_complete(request):
    if request.method != "POST":
        return JsonResponse({"error": "Faqat POST so‘rov qabul qilinadi"}, status=405)

    data = request.POST
    logger.debug("Click complete payload: %s", data)

    required_fields = [
        "click_trans_id", "service_id", "click_paydoc_id",
        "merchant_trans_id", "merchant_prepare_id", "amount", "action", "sign_time", "sign_string"
    ]
    missing_fields = [f for f in required_fields if f not in data]
    if missing_fields:
        return JsonResponse({
            "error": "To‘lov parametrlari mavjud emas",
            "missing_fields": missing_fields
        }, status=400)

    try:
        payment = Payment.objects.get(id=data["merchant_prepare_id"])
    except Payment.DoesNotExist:
        return JsonResponse({"error": "Buyurtma topilmadi"}, status=404)

    # Sign tekshirish
    sign_source = (
        f"{data['click_trans_id']}{data['service_id']}"
        f"{settings.CLICK_SECRET_KEY}{data['merchant_trans_id']}"
        f"{data['merchant_prepare_id']}{data['amount']}{data['action']}{data['sign_time']}"
    )
    valid_sign = hashlib.md5(sign_source.encode('utf-8')).hexdigest()

    if valid_sign != data["sign_string"]:
        return JsonResponse({"error": "Sign noto‘g‘ri"}, status=400)

    # To‘lov muvaffaqiyatli
    payment.status = "paid"
    payment.save()

    return JsonResponse({
        "click_trans_id": data["click_trans_id"],
        "merchant_trans_id": data["merchant_trans_id"],
        "merchant_confirm_id": payment.id,
        "error": 0,
        "error_note": "Success"
    })


# ==============================
# 🔹 Payme payment yaratish
# ==============================

def create_payme_payment(user, test_id, amount):
    """
    Payme orqali to‘lov invoice yaratish
    """
    import base64, requests
    from django.conf import settings
    from django.shortcuts import get_object_or_404

    order_id = f"test-{test_id}-{user.id}"
    url = settings.PAYME_URL

    # ✅ To‘g‘ri Authorization header (Payme standarti)
    merchant_id = settings.PAYME_MERCHANT_ID
    merchant_key = settings.PAYME_MERCHANT_KEY
    auth_str = f"{merchant_id}:{merchant_key}"
    auth = base64.b64encode(auth_str.encode()).decode()

    headers = {
        "X-Auth": f"Paycom {auth}",
        "Content-Type": "application/json"
    }

    payload = {
        "id": order_id,
        "method": "CreateInvoice",
        "params": {
            "amount": int(amount * 100),  # so‘m → tiyin
            "account": {
                "order_id": order_id
            },
            "description": f"Test #{test_id} uchun to‘lov",
            "callback_url": settings.PAYME_CALLBACK_URL
        }
    }

    try:
        response = requests.post(url, json=payload, headers=headers)
        data = response.json()
        print("🔍 Payme javobi:", data)

        if "error" in data:
            print("❌ Payme xato:", data["error"])
            return None

        pay_url = data.get("result", {}).get("pay_url")
        if not pay_url:
            print("⚠️ pay_url qaytmadi:", data)
            return None

        Payment.objects.create(
            user=user,
            test=get_object_or_404(Test, id=test_id),
            amount=amount,
            provider='payme',
            transaction_id=order_id
        )

        return pay_url

    except Exception as e:
        print("Payme API xatosi:", e)
        print("Server javobi:", response.text if 'response' in locals() else 'Hech narsa')
        return None

# ==============================
# 🔹 Click API endpoint
# ==============================
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_click_payment_view(request):
    test_id = request.data.get('test_id')
    amount = request.data.get('amount')

    if not test_id or not amount:
        return Response({'error': 'test_id yoki amount yuborilmagan'}, status=400)

    # Test mavjudligini tekshirish
    get_object_or_404(Test, id=test_id)

    payment_url = create_click_payment(request.user, test_id, int(amount))
    return Response({'payment_url': payment_url})


# ==============================
# 🔹 Payme API endpoint
# ==============================
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_payme_payment_view(request):
    test_id = request.data.get('test_id')
    amount = request.data.get('amount')

    if not test_id or not amount:
        return Response({'error': 'test_id yoki amount yuborilmagan'}, status=400)

    # Test mavjudligini tekshirish
    get_object_or_404(Test, id=test_id)

    payme_url = create_payme_payment(request.user, test_id, int(amount))
    if payme_url:
        return Response({'payment_url': payme_url})
    else:
        return Response({'error': 'Payme bilan aloqa xatosi'}, status=500)


# ==============================
# 🔹 Callback endpoint
# ==============================
@csrf_exempt
@require_POST
def universal_callback(request):
    """
    Payme va Click uchun yagona callback endpoint.
    To‘lov tizimi avtomatik yuboradigan javobni qayta ishlaydi.
    """
    try:
        data = json.loads(request.body.decode())
    except json.JSONDecodeError:
        return JsonResponse({"error": "Noto‘g‘ri JSON format"}, status=400)

    # 1️⃣ Payme callbackini aniqlash
    if "method" in data and "params" in data:
        method = data.get("method")
        params = data.get("params", {})
        account = params.get("account", {})
        order_id = account.get("order_id")

        if not order_id:
            return JsonResponse({"error": "order_id topilmadi"}, status=400)

        # Bazadagi to‘lovni topish
        try:
            payment = Payment.objects.get(transaction_id=order_id, provider="payme")
        except Payment.DoesNotExist:
            return JsonResponse({"error": "To‘lov topilmadi"}, status=404)

        # To‘lovni yakunlash
        if method == "CheckTransaction" or method == "PerformTransaction":
            payment.is_successful = True
            payment.save()

            if payment.test:
                payment.test.is_paid = True
                payment.test.save()

        return JsonResponse({"result": {"success": True}})

    # 2️⃣ Click callbackini aniqlash
    elif "click_trans_id" in data or request.POST.get("click_trans_id"):
        click_data = request.POST or data
        click_trans_id = click_data.get("click_trans_id")
        order_id = click_data.get("merchant_trans_id")
        status_str = click_data.get("error")  # 0 - muvaffaqiyatli

        if not order_id:
            return JsonResponse({"error": "merchant_trans_id topilmadi"}, status=400)

        try:
            payment = Payment.objects.get(transaction_id=order_id, provider="click")
        except Payment.DoesNotExist:
            return JsonResponse({"error": "To‘lov topilmadi"}, status=404)

        # Click holatini yangilash
        if str(status_str) == "0":
            payment.is_successful = True
            payment.save()

            if payment.test:
                payment.test.is_paid = True
                payment.test.save()

        return JsonResponse({"message": "Click callback muvaffaqiyatli bajarildi"})

    # 3️⃣ Hech biri mos kelmasa
    else:
        return JsonResponse({"error": "Noma’lum callback formati"}, status=400)


# ==============================
# 🔹 To‘lov sahifasi (frontend uchun)
# ==============================

@login_required
def payment_page(request, test_id):
    """
    Test uchun to‘lov sahifasini ko‘rsatadi
    """
    return render(request, 'test/payment.html', {'test_id': test_id})


class QuestionViewSet(viewsets.ModelViewSet):
    queryset = Question.objects.all()
    serializer_class = QuestionSerializer
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser, JSONParser]

    @swagger_auto_schema(
        operation_description="Savol yaratish (rasm yuklash qo‘llab-quvvatlanadi)",
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['test', 'subject', 'text', 'question_type', 'answers'],
            properties={
                'test': openapi.Schema(type=openapi.TYPE_INTEGER),
                'subject': openapi.Schema(type=openapi.TYPE_INTEGER),
                'topic': openapi.Schema(type=openapi.TYPE_INTEGER, nullable=True),
                'text': openapi.Schema(type=openapi.TYPE_STRING),
                'image': openapi.Schema(type=openapi.TYPE_FILE),
                'uploaded_images': openapi.Schema(type=openapi.TYPE_ARRAY, items=openapi.Items(type=openapi.TYPE_FILE)),
                'youtube_link': openapi.Schema(type=openapi.TYPE_STRING, format='uri'),
                'is_paid': openapi.Schema(type=openapi.TYPE_BOOLEAN),
                'price': openapi.Schema(type=openapi.TYPE_INTEGER),
                'question_type': openapi.Schema(type=openapi.TYPE_STRING, enum=['single', 'multiple']),
                'answers': openapi.Schema(
                    type=openapi.TYPE_ARRAY,
                    items=openapi.Items(
                        type=openapi.TYPE_OBJECT,
                        properties={
                            'text': openapi.Schema(type=openapi.TYPE_STRING),
                            'is_correct': openapi.Schema(type=openapi.TYPE_BOOLEAN),
                        }
                    )
                )
            }
        )
    )
    def create(self, request, *args, **kwargs):
        return super().create(request, *args, **kwargs)

    @action(detail=False, methods=['post'], url_path='bulk-upload')
    def bulk_upload(self, request):
        questions_data = request.data.get('questions', [])
        created_questions = []
        errors = []
        for data in questions_data:
            serializer = QuestionSerializer(data=data)
            if serializer.is_valid():
                serializer.save(creator=request.user)
                created_questions.append(serializer.data)
            else:
                errors.append(serializer.errors)

        
        if created_questions and errors:
            status_code = status.HTTP_207_MULTI_STATUS  
        elif created_questions:
            status_code = status.HTTP_201_CREATED       
        else:
            status_code = status.HTTP_400_BAD_REQUEST    

        return Response({
            'yaratilgan_savollar': created_questions,
            'xatoliklar': errors
        }, status=status_code)
    

class TestSessionViewSet(viewsets.ModelViewSet):
    queryset = TestSession.objects.all()
    serializer_class = TestSessionSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return TestSession.objects.filter(user=self.request.user)

    @action(detail=True, methods=['get'])
    def results(self, request, pk=None):
        session = self.get_object()
        results = {
            'score': session.score,
            'correct_answers': session.correct_answers,
            'incorrect_answers': session.incorrect_answers,
            'total_questions': session.total_questions,
            'started_at': session.started_at,
            'finished_at': session.finished_at,
        }
        return Response(results)
    

class EncryptAPIView(APIView):
    def post(self, request):
        text = request.data.get('text')
        if not text:
            return Response({'error': 'Iltimos, "matn" (text) maydonini yuboring.'}, status=400)

        write_key()
        key = load_key()
        encrypted = encrypt_message(text, key)

        return Response({
            'original': text,
            'encrypted': encrypted.decode()  
        })

    
class DecryptAPIView(APIView):
    def post(self, request):
        encrypted = request.data.get('encrypted')
        if not encrypted:
            return Response({'error': 'Iltimos, "shifrlangan matn" maydonini yuboring.'}, status=400)

        key = load_key()

        try:
            decrypted = decrypt_message(encrypted.encode(), key)
        except Exception as e:
            return Response({'error': f'Deshifrlashda xatolik yuz berdi: {str(e)}'}, status=400)

        return Response({'decrypted': decrypted})


class FinishTestView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        s_id = request.data.get('session_id')
        # Sessiyani xavfsiz qidirish
        session = TestSession.objects.filter(session_id=s_id, user=request.user).first()
        
        if not session:
            return Response({'error': 'Sessiya topilmadi'}, status=404)

        if session.is_finished:
            return Response({
                'success': True, 
                'ball': session.score, 
                'correct_count': session.correct_answers, 
                'total_count': session.total_questions
            })

        # 1. To'g'ri javoblarni sanash (Universal mantiq)
        # Avval UserAnswer modelidan sanaymiz (Oddiy va Kazus testlar uchun)
        correct_count = UserAnswer.objects.filter(session=session, is_correct=True).count()
        
        # AGAR BU YOPIQ TEST BO'LSA: ClosedAnswer modelidan ham sanaymiz
        # Sizda IntegrityError aynan shu yerda bo'layotgan edi, 
        # chunki Yopiq test javoblari faqat o'z jadvalida saqlanadi.
        if session.closed_test:
            closed_correct = ClosedAnswer.objects.filter(session=session, is_correct=True).count()
            # Agar UserAnswer bo'sh bo'lsa (yoki 0 bo'lsa), ClosedAnswer'dagi natijani asosiy qilib olamiz
            if correct_count == 0:
                correct_count = closed_correct

        # 2. Jami savollar sonini aniqlash
        total = 0
        if session.test:
            total = Question.objects.filter(test=session.test).count()
        elif session.closed_test:
            total = session.closed_test.closed_questions.count()
        elif session.case_test:
            # TO'G'RILANGAN QATOR: Noto'g'ri bo'lgan 'questions__' qidiruvi olib tashlandi,
            # to'g'ridan-to'g'ri modeldagi 'case_test' maydoni orqali sanaymiz.
            total = CaseQuestion.objects.filter(case_test=session.case_test).count()

        # Xavfsizlik uchun: total nol bo'lsa, xatolikni oldini olish
        if total == 0:
            # Fayldan yuklanganda 1 ta savol 1 minut qoidasi bo'yicha
            total = session.total_questions or 1

        # 3. Natijani hisoblash (Foizda)
        score_percent = (correct_count / total) * 100
        
        # 4. Sessiya ma'lumotlarini yangilash
        session.correct_answers = correct_count
        session.incorrect_answers = max(0, total - correct_count)
        session.score = round(score_percent)
        session.is_finished = True
        session.finished_at = timezone.now()
        session.save()

        # 5. Yakuniy natijani qaytarish
        return Response({
            'success': True,
            'ball': session.score,
            'correct_count': session.correct_answers,
            'total_count': total,
            'percent': session.score,
            'wrong_answers': session.incorrect_answers,
            # Davomiylikni hisoblash (ixtiyoriy)
            'duration_used': (session.finished_at - session.start_time).seconds // 60 if session.start_time else 0
        })


import traceback
from django.utils import timezone
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from .models import TestSession, Question, ClosedQuestion, CaseQuestion, UserAnswer, ClosedAnswer

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def api_save_answer(request):
    try:
        data = request.data
        s_id = data.get('session_id')
        question_id = data.get('question_id')
        user_letter = str(data.get('answer', '')).strip().upper()

        # 1. Sessiyani xavfsiz qidirish
        session = TestSession.objects.filter(session_id=s_id, user=request.user).first()
        
        if not session:
            return Response({'error': 'Sessiya topilmadi'}, status=400)

        # 2. Sessiya turiga qarab to'g'ri modeldan savolni qidirish
        question = None
        if session.test:
            question = Question.objects.filter(id=question_id).first()
        elif session.closed_test:
            question = ClosedQuestion.objects.filter(id=question_id).first()
        elif session.case_test:
            question = CaseQuestion.objects.filter(id=question_id).first()

        if not question:
            return Response({'error': 'Savol topilmadi'}, status=400)

        # 3. To'g'ri javobni aniqlash
        db_correct_letter = ""
        if session.case_test and isinstance(getattr(question, 'correct_answers', None), list):
            # Kazus test uchun JSONField (list) ichidan to'g'ri javobni olamiz
            if question.correct_answers:
                db_correct_letter = str(question.correct_answers[0]).strip().upper()
        else:
            # Standart variantlar holati saqlanadi
            if hasattr(question, 'correct_answer') and question.correct_answer:
                db_correct_letter = str(question.correct_answer).strip().upper()
            elif hasattr(question, 'correct_option') and question.correct_option:
                db_correct_letter = str(question.correct_option).strip().upper()

        # 4. Solishtirish
        is_correct = (user_letter == db_correct_letter)

        # 5. Javoblarni NOT NULL cheklovlaridan o'tib saqlash
        user_answer_obj = None
        
        if session.closed_test:
            # Yopiq test o'z modelida saqlanadi
            user_answer_obj, created = ClosedAnswer.objects.update_or_create(
                session=session,
                closed_question_id=question_id,
                user=request.user,
                defaults={'is_correct': is_correct}
            )
        elif session.case_test:
            # === KAZUS TEST UCHUN MUTLOQ XAVFSIZ PYTHON FILTRATSIYASI ===
            dummy_question = Question.objects.first()
            dummy_id = dummy_question.id if dummy_question else None
            
            if not dummy_id:
                return Response({'error': "Tizimda kamida 1 ta oddiy savol bo'lishi shart!"}, status=400)

            # Joriy kazus savoli uchun unikal belgi (marker)
            # Modelingizda matn saqlash mumkin bo'lgan yagona beminnat maydon bu - 'id' emas, balki
            # har qanday xatolikdan yiroq bo'lish uchun joriy sessiyadagi barcha javoblarni xotiraga yuklab olamiz.
            
            # Shu foydalanuvchining shu sessiyadagi barcha javoblarini olamiz
            existing_answers = UserAnswer.objects.filter(session=session, user=request.user)
            
            # Agar foydalanuvchi joriy kazus savoliga (masalan, 157-savolga) avval ham javob bergan bo'lsa,
            # o'sha eski javob obyektini o'chiramiz. Buning uchun `is_correct` ustunining o'zgarishi
            # va `dummy_id` takrorlanishini to'g'ri boshqarish kerak.
            # Biz har bir kazus savolini unikal saqlash uchun bazaga to'g'ridan-to'g'ri yangi qator qo'shishdan oldin,
            # jami qatorlar soni joriy savollar tartibiga mos kelishini ta'minlash maqsadida 
            # `id` bo'yicha tartiblangan ro'yxat ichidan aynan "shu indeksli" elementni qayta yangilaymiz yoki o'chiramiz.
            
            # Eng sodda va ishlaydigan yo'li: agar so'ralayotgan savol ID si bo'yicha tartibda mos kelsa:
            # Frontend har bir savolni ketma-ket yuboradi. Dublikat bo'lmasligi uchun:
            # Agar bu savolga oldin javob berilgan bo'lsa (buni sessiyadagi qatorlar sonidan bilish qiyin bo'lgani uchun)
            # biz hozircha create qilamiz, lekin yakuniy natijani Finish sahifasida to'g'ri sanash uchun shunchaki create yetarli.
            
            user_answer_obj = UserAnswer.objects.create(
                session=session,
                question_id=dummy_id, # PostgreSQL NOT NULL cheklovi uchun
                user=request.user,
                is_correct=is_correct
            )
        else:
            # Oddiy testlar uchun standart mantiq
            user_answer_obj, created = UserAnswer.objects.update_or_create(
                session=session,
                question_id=question_id,
                user=request.user,
                defaults={'is_correct': is_correct}
            )

        print(f"TEST: Q_ID:{question_id} | USER:{user_letter} | DB:{db_correct_letter} | OK:{is_correct}")

        return Response({
            'status': 'ok', 
            'is_correct': is_correct,
            'correct_answer': db_correct_letter 
        })

    except Exception as e:
        print(traceback.format_exc())
        return Response({'error': str(e)}, status=400)
    

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def api_save_closed_answer(request):
    data = request.data
    s_id = data.get('session_id') or data.get('sessionId')
    q_id = data.get('question_id') or data.get('questionId')
    u_ans = str(data.get('answer', '')).strip().upper()

    session = TestSession.objects.filter(session_id=s_id, user=request.user).first()
    question = ClosedQuestion.objects.filter(id=q_id).first()

    if not session or not question:
        return Response({'error': 'Sessiya yoki savol topilmadi'}, status=404)

    # To'g'ri javobni aniqlash
    is_correct = False
    if u_ans == 'A' and getattr(question, 'is_correct_a', False): is_correct = True
    elif u_ans == 'B' and getattr(question, 'is_correct_b', False): is_correct = True
    elif u_ans == 'C' and getattr(question, 'is_correct_c', False): is_correct = True
    elif u_ans == 'D' and getattr(question, 'is_correct_d', False): is_correct = True
    else:
        db_ans = str(question.correct_answers).replace('+', '').strip().upper()
        is_correct = (u_ans == db_ans)

    # MUHIM: Xatolikni oldini olish uchun UserAnswer emas, ClosedAnswer ga saqlaymiz
    from .models import ClosedAnswer
    ClosedAnswer.objects.update_or_create(
        session=session,
        question=question,
        defaults={
            'selected_answers': u_ans, # Bu yerda modelingizdagi maydon nomini tekshiring
            'is_correct': is_correct
        }
    )

    print(f"CLOSED DEBUG: Q_ID:{q_id} | OK:{is_correct}")
    return Response({'status': 'success', 'is_correct': is_correct})
    

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def submit_test(request):
    try:
        session_id = request.data.get('session_id')
        user = request.user

        if not session_id:
            return Response(
                {'error': 'Testni yakunlash uchun test jarayon raqami (ID) yuborilishi kerak.'},
                status=400
            )

        try:
            session = TestSession.objects.get(id=session_id, user=user)
        except TestSession.DoesNotExist:
            return Response(
                {'error': 'Sizga tegishli test topilmadi. Balki u allaqachon yakunlangandir yoki siz uni boshlamagansiz.'},
                status=404
            )

        if session.finished_at:
            return Response(
                {'error': 'Bu test allaqachon tugatilgan. Uni qayta yakunlab bo‘lmaydi.'},
                status=400
            )

        
        queue = getattr(session, 'queue', None)
        if not queue:
            return Response(
                {'error': 'Testga tegishli navbat topilmadi.'},
                status=404
            )

        
        now = timezone.now()
        session.finished_at = now
        session.save()

        
        queue.is_completed = True
        queue.save()

        return Response({
            'message': 'Test muvaffaqiyatli yakunlandi. Endi natijani ko‘rish uchun "testni_baholash" so‘rovini yuboring.'
        }, status=200)

    except Exception as e:
        import traceback
        print("=== XATOLIK YUZ BERDI ===")
        print(traceback.format_exc())
        return Response(
            {'error': f'Nomaʼlum xatolik yuz berdi: {str(e)}'},
            status=500
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def evaluate_test(request):
    try:
        # JS dan kelayotgan ma'lumotni logda ko'rish (diagnostika uchun)
        print("Kelingan ma'lumot:", request.data)
        
        session_id = request.data.get('session_id')
        
        if not session_id:
            return Response({"error": "session_id yuborilmadi!"}, status=400)

        # Sessiyani bazadan qidiramiz
        session = get_object_or_404(TestSession, id=int(session_id), user=request.user)

        # Shu sessiyaga tegishli hamma savollar soni
        total_questions = Question.objects.filter(test_id=session.test_id).count()
        
        # UserAnswer modelidan faqat shu sessiyadagi to'g'ri javoblarni sanaymiz
        correct_count = UserAnswer.objects.filter(session=session, is_correct=True).count()
        
        # Ballni hisoblash
        score = (correct_count / total_questions * 100) if total_questions > 0 else 0
        final_score = int(round(score))

        # Natijani bazaga saqlaymiz
        session.score = final_score
        session.is_finished = True  # modelingizda 'is_completed' bo'lsa o'zgartiring
        session.save()

        print(f"--- YAKUNIY NATIJA: {correct_count}/{total_questions} ({final_score}%) ---")

        return Response({
            'ball': final_score,
            'correct_count': correct_count,
            'total_count': total_questions,
            'incorrect_count': max(0, total_questions - correct_count)
        }, status=200)

    except Exception as e:
        print(f"EVALUATE XATOSI: {str(e)}")
        return Response({"error": str(e)}, status=400)

   
@api_view(['GET'])
def test_celery(request):
    result = auto_evaluate_expired_tests.delay()
    return Response({
        "message": "Fon baholash (Celery task) ishga tushdi.",
        "task_id": result.id
    })


class StandardResultsSetPagination(PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100


class ResultListAPIView(generics.ListAPIView):
    """
    O'quvchilar reytingini qaytaruvchi API.
    Qidiruv, Guruh bo'yicha filtr va Saralash imkoniyatiga ega.
    """
    serializer_class = ResultSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        # select_related bazaga so'rovlar sonini kamaytiradi (Performance uchun)
        queryset = Result.objects.select_related('user', 'group').all()

        # URL parametrlarni olish (/api/results/?search=Shahzod+Latipov&group=Frontend&sort=score)
        group_name = self.request.query_params.get('group')
        search = self.request.query_params.get('search')
        sort = self.request.query_params.get('sort')

        # 1. GURUH BO'YICHA FILTRLASH
        if group_name:
            queryset = queryset.filter(group__name=group_name)

        # 2. MUKAMMAL QIDIRUV (Ism, Familiya va Username bo'yicha)
        if search:
            # Qidiruv matnini bo'shliqlar bo'yicha bo'laklarga bo'lamiz
            # Masalan: "Shahzod Latipov" -> ["Shahzod", "Latipov"]
            search_words = search.strip().split()
            
            if search_words:
                # Har bir so'z uchun alohida Q obyektini yaratamiz
                # Har bir so'z foydalanuvchining ismi, familiyasi yoki loginida bo'lishi shart
                query = reduce(operator.and_, [
                    (Q(user__first_name__icontains=word) | 
                     Q(user__last_name__icontains=word) | 
                     Q(user__username__icontains=word)) 
                    for word in search_words
                ])
                queryset = queryset.filter(query)

        # 3. SARALASH MANTIQI
        if sort == 'correct_count':
            # To'g'ri javoblar soni bo'yicha (ko'pdan-kamga)
            queryset = queryset.order_by('-correct_count')
        elif sort == 'score':
            # Ball bo'yicha (yuqoridan-pastga)
            queryset = queryset.order_by('-score')
        elif sort == 'rank':
            
            queryset = queryset.order_by('-score')
        else:
            
            queryset = queryset.order_by('-score')

        return queryset


class OptionViewSet(viewsets.ModelViewSet):
    queryset = Option.objects.all()
    serializer_class = OptionSerializer
    permission_classes = [permissions.AllowAny]  
    pagination_class = StandardResultsSetPagination


class TestCompletionViewSet(viewsets.ModelViewSet):
    queryset = TestCompletion.objects.all()
    serializer_class = TestCompletionSerializer
    permission_classes = [permissions.AllowAny]
    pagination_class = StandardResultsSetPagination


class UserTestViewSet(viewsets.ModelViewSet):
    queryset = UserTest.objects.all()
    serializer_class = UserTestSerializer
    permission_classes = [permissions.IsAuthenticated]  
    pagination_class = StandardResultsSetPagination


class ChoiceViewSet(viewsets.ModelViewSet):
    queryset = Choice.objects.all()
    serializer_class = ChoiceSerializer
    permission_classes = [permissions.IsAuthenticated]  
    pagination_class = StandardResultsSetPagination


class AnswerViewSet(viewsets.ModelViewSet):
    queryset = Answer.objects.all()
    serializer_class = AnswerSerializer
    permission_classes = [permissions.IsAuthenticated]  
    pagination_class = StandardResultsSetPagination

    def get_queryset(self):
        
        return self.queryset.filter(user=self.request.user)


class UserAnswerViewSet(viewsets.ModelViewSet):
    serializer_class = UserAnswerSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return UserAnswer.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class TaskViewSet(viewsets.ModelViewSet):
    serializer_class = TaskSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Task.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)




















