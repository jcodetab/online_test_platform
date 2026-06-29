from cryptography.fernet import Fernet
from .models import Test, TestAnswer, ClosedTest, CaseTest, CaseQuestion, Question, ClosedQuestion
import os
import docx
from docx import Document
from django.db import transaction
import re
import io





KEY_FILE = os.path.join(os.path.dirname(__file__), 'secret.key')


def write_key():
    """
    Yangi kalit generatsiya qilinadi va secret.key fayliga yoziladi,
    agar u mavjud bo'lmasa.
    """
    if not os.path.exists(KEY_FILE):
        key = Fernet.generate_key()
        try:
            with open(KEY_FILE, 'wb') as f:
                f.write(key)
        except IOError as e:
            raise Exception(f"Kalit yozishda xatolik yuz berdi: {e}")


def load_key():
    """
    Kalit fayldan o'qiladi. Agar mavjud bo'lmasa, xatolik chiqariladi.
    """
    if not os.path.exists(KEY_FILE):
        raise FileNotFoundError("Kalit fayli topilmadi. Avval write_key() funksiyasini chaqiring.")
    
    try:
        with open(KEY_FILE, 'rb') as f:
            return f.read()
    except IOError as e:
        raise Exception(f"Kalit o'qishda xatolik: {e}")


def encrypt_message(message, key):
    """
    Matnni shifrlaydi (Fernet orqali).
    """
    if not isinstance(message, str):
        raise ValueError("Faqat string (matn) shifrlanadi.")
    
    f = Fernet(key)
    return f.encrypt(message.encode())


def decrypt_message(encrypted_message, key):
    """
    Shifrlangan matnni deshifrlaydi.
    """
    f = Fernet(key)
    try:
        return f.decrypt(encrypted_message).decode()
    except Exception as e:
        raise Exception(f"Deshifrlashda xatolik: {e}")
    

def evaluate_test(session):
    answers = TestAnswer.objects.filter(session=session)
    correct = 0
    for answer in answers:
        if set(answer.selected_answers) == set(answer.question.correct_answers):
            correct += 1
    score = round((correct / answers.count()) * 100, 2) if answers.exists() else 0.0
    session.queue.score = score
    session.queue.is_completed = True
    session.queue.save()
    session.is_finished = True
    session.auto_evaluated = True
    session.save()


def parse_and_create_case_test(file, user, subject_id=None):
    """
    Word (.docx) faylidan kazus matni va savollarni o'qib,
    CaseTest va CaseQuestion modellariga saqlaydi.
    """
    doc = docx.Document(file)
    full_text = []
    
    # Barcha matnni yig'ib olamiz
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    if not full_text:
        return None

    # 1. Kazus matnini aniqlash
    # Odatda birinchi xatboshi kazus matni (muammo bayoni) bo'ladi
    title = file.name
    case_description = full_text[0] # Birinchi xatboshini kazus matni deb olamiz
    
    # 2. CaseTest ob'ektini yaratish
    new_case_test = CaseTest.objects.create(
        creator=user,
        title=title,
        case_text=case_description,
        subject_id=subject_id,
        file=file
    )

    # 3. Savollarni ajratish va yaratish
    # Bu qismda matndan savol va variantlarni qidiramiz (regex yoki oddiy mantiq)
    current_question = None
    
    # 1-xatboshidan keyingi hamma narsani savol deb qaraymiz
    for line in full_text[1:]:
        # Agar satr raqam bilan boshlansa (masalan: "1. Savol...")
        if line[0].isdigit() and ('.' in line[:3] or ')' in line[:3]):
            current_question = CaseQuestion.objects.create(
                case_test=new_case_test,
                text=line,
                answer_a="", # Keyingi qatorlarda to'ldiriladi
                answer_b="",
                correct_answers=[]
            )
        # Agar variantlar bo'lsa (A, B, C, D)
        elif current_question:
            clean_line = line.strip()
            if clean_line.startswith(('A)', 'A.', 'a)', 'a.')):
                current_question.answer_a = clean_line
            elif clean_line.startswith(('B)', 'B.', 'b)', 'b.')):
                current_question.answer_b = clean_line
            elif clean_line.startswith(('C)', 'C.', 'c)', 'c.')):
                current_question.answer_c = clean_line
            elif clean_line.startswith(('D)', 'D.', 'd)', 'd.')):
                current_question.answer_d = clean_line
            
            # To'g'ri javobni belgilash (agar Word'da belgi bo'lsa, masalan: # yoki +)
            if '#' in clean_line or '+' in clean_line:
                # To'g'ri javobni JSONField ga saqlaymiz
                variant_letter = clean_line[0].upper()
                if variant_letter not in current_question.correct_answers:
                    current_question.correct_answers.append(variant_letter)
            
            current_question.save()

    return new_case_test
    

def parse_full_docx(file):
    doc = docx.Document(file)
    questions = []
    current_q = None
    
    for para in doc.paragraphs:
        # \xa0 (non-breaking space) larni oddiy probelga aylantiramiz
        text = para.text.strip().replace('\xa0', ' ')
        if not text: continue

        if text.lower().startswith('[s]'):
            if current_q: questions.append(current_q)
            current_q = {'text': text[3:].strip(), 'ans': [], 'corr_idx': 0}
            option_count = 0
        
        elif current_q and text[0] in ['+', '-']:
            if len(current_q['ans']) < 4:
                # Birinchi belgidan keyin hamma narsani olamiz va tozalaymiz
                content = text[1:].strip()
                
                # HTML qavslarini xavfsiz saqlash (escape qilish shart emas, lekin bo'sh bo'lmasligi kerak)
                if content:
                    current_q['ans'].append(content)
                    if text[0] == '+':
                        current_q['corr_idx'] = len(current_q['ans']) - 1
    
    if current_q: questions.append(current_q)
    return questions


def parse_and_create_questions(test_instance, docx_file, user):
    """
    Savollarni Word fayldan o'qiydi va tegishli modelga (Test, CaseTest, ClosedTest) saqlaydi.
    """
    try:
        # Modellarni dinamik import qilish
        from .models import Question, CaseQuestion, ClosedQuestion, Test, CaseTest, ClosedTest

        docx_file.seek(0)
        doc = Document(io.BytesIO(docx_file.read()))
        
        # 1. 🎯 FAYL NOMIDAN TURINI 100% ANIQLASH (Eng ishonchli usul)
        file_name_lower = str(getattr(docx_file, 'name', '')).lower()
        model_name = test_instance.__class__.__name__
        instance_type = getattr(test_instance, 'test_type', '').lower()

        # Front-end yoki baza nima deyishidan qat'i nazar, fayl nomida shu so'zlar bo'lsa - BU KAZUS!
        is_case_file = 'kazus' in file_name_lower or 'kasuz' in file_name_lower or 'case' in file_name_lower
        
        is_case = model_name == 'CaseTest' or instance_type == 'kazus' or is_case_file
        is_closed = model_name == 'ClosedTest' or instance_type == 'closed'
        is_regular = not is_case and not is_closed

        print(f"--- 🚨 PARSER: Fayl nomi = {file_name_lower} ---")
        print(f"--- 🚨 PARSER: IsCase = {is_case} (File trigger: {is_case_file}) ---")

        # 🎯 AGAR KAZUS BO'LSA VA OBYEKT ADASHIB 'TEST' (ODDIY) BO'LIB KELGAN BO'LSA, UNI TO'G'RILAYMIZ:
        if is_case and model_name == 'Test':
            print("--- 🎯 PARSER: Ota model aniqlandi, uni KAZUS formatiga majburiy o'tkazamiz! ---")
            test_instance.test_type = 'kazus'
            test_instance.save(update_fields=['test_type'])

        questions = []
        current_q = None

        # MATNNI FILTRLASH VA SAVOLLARNI AJRATISH
        for para in doc.paragraphs:
            text = para.text.strip().replace('\xa0', ' ').replace('\u200b', '')
            if not text: 
                continue

            is_explicit_q = re.match(r'^\[s\]', text, re.IGNORECASE)
            is_question_mark = text.endswith('?')
            is_variant = re.match(r'^[+\-–—]', text)

            if is_explicit_q or is_question_mark:
                if current_q and current_q['ans']:
                    questions.append(current_q)
                
                q_text = re.sub(r'^\[s\]\s*', '', text, flags=re.IGNORECASE).strip() if is_explicit_q else text
                current_q = {'text': q_text, 'ans': [], 'corr_idx': 0}

            elif current_q is not None and is_variant:
                sign = text[0]
                val = text[1:].strip()
                if val:
                    current_q['ans'].append(val)
                    if sign == '+':
                        current_q['corr_idx'] = len(current_q['ans']) - 1

            elif not is_variant:
                if current_q is None:
                    current_q = {'text': text, 'ans': [], 'corr_idx': 0}
                else:
                    if not current_q['ans']:
                        current_q['text'] += "\n" + text

        if current_q and current_q['ans']:
            questions.append(current_q)

        if not questions: 
            print("--- PARSER: Fayldan hech qanday savol topilmadi! ---")
            return 0

        # BAZAGA SAQLASH
        db_objs = []
        char_map = ['A', 'B', 'C', 'D']
        
        with transaction.atomic():
            for q_data in questions:
                ans = q_data['ans']
                while len(ans) < 4: 
                    ans.append("---")
                
                idx = q_data['corr_idx']
                correct_letter = char_map[idx] if idx < 4 else 'A'

                if is_case:
                    # 🎯 KAZUS MODELINI TEKSHIRISH VA ULASH
                    target_case_instance = test_instance
                    if model_name == 'Test':
                        target_case_instance = CaseTest.objects.filter(parent_test=test_instance).first()
                        if not target_case_instance:
                            target_case_instance = CaseTest.objects.create(
                                parent_test=test_instance,
                                title=test_instance.title,
                                creator=user,
                                subject=test_instance.subject,
                                case_text=test_instance.title,
                                creation_method='file'
                            )

                    db_objs.append(CaseQuestion(
                        case_test=target_case_instance, 
                        question=q_data['text'], 
                        answer_a=ans[0], 
                        answer_b=ans[1], 
                        answer_c=ans[2], 
                        answer_d=ans[3],
                        correct_answers=[correct_letter],
                        question_type='case'
                    ))
                elif is_closed:
                    db_objs.append(ClosedQuestion(
                        test=test_instance, 
                        text=q_data['text'],
                        answer_A=ans[0], 
                        answer_B=ans[1], 
                        answer_C=ans[2], 
                        answer_D=ans[3],
                        correct_answers=correct_letter, 
                        question_type='single'
                    ))
                elif is_regular:
                    db_objs.append(Question(
                        test=test_instance, 
                        question=q_data['text'], 
                        answer_a=ans[0], 
                        answer_b=ans[1], 
                        answer_c=ans[2], 
                        answer_d=ans[3],
                        correct_option=correct_letter, 
                        creator=user,
                        is_correct_a=(idx==0), 
                        is_correct_b=(idx==1), 
                        is_correct_c=(idx==2), 
                        is_correct_d=(idx==3)
                    ))

            # Bulk create qilish
            if is_case and db_objs: 
                CaseQuestion.objects.bulk_create(db_objs)
                print(f"--- PARSER: {len(db_objs)} ta CaseQuestion saqlandi. ---")
            elif is_closed and db_objs: 
                ClosedQuestion.objects.bulk_create(db_objs)
            elif is_regular and db_objs: 
                Question.objects.bulk_create(db_objs)

            # VAQTNI VA TEST TURINI ANIQ YANGILASH
            created_count = len(db_objs)
            if created_count > 0:
                for field in ['test_duration', 'duration', 'duration_minutes']:
                    if hasattr(test_instance, field):
                        setattr(test_instance, field, created_count)
                
                if is_case:
                    if hasattr(test_instance, 'test_type'):
                        test_instance.test_type = 'kazus'
                    
                    if hasattr(test_instance, 'parent_test') and test_instance.parent_test:
                        parent = test_instance.parent_test
                        parent.test_duration = created_count
                        parent.test_type = 'kazus'
                        parent.save(update_fields=['test_duration', 'test_type'])
                
                test_instance.save()

            return created_count

    except Exception as e:
        print(f"--- PARSING ERROR: {str(e)} ---")
        import traceback
        traceback.print_exc()
        return 0

