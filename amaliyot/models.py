from django.db import models
from django.contrib.auth import get_user_model
from django.utils import timezone
import uuid
import string
import random
from django.contrib.auth.models import User
from django.db.models.signals import post_save, post_delete
from django.dispatch import receiver
from decimal import Decimal
from django.conf import settings
from django.utils.crypto import get_random_string
import io
import re
from docx import Document






User = get_user_model()

def generate_random_key(length=10):
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))


class ClosedTest(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='closed_tests', null=True, blank=True)
    subject = models.ForeignKey('Subject', on_delete=models.SET_NULL, null=True, blank=True, related_name='closed_tests')
    title = models.CharField(max_length=255)
    description = models.TextField(blank=True)
    
    # default=0 qoladi. Admin panelda 0 kiritilsa - savol soni olinadi, 
    # aks holda kiritilgan raqam saqlanadi.
    duration = models.IntegerField(default=0, help_text="Vaqt minutda (0 bo'lsa savollar soni olinadi)") 
    
    score = models.IntegerField(default=100)
    file = models.FileField(upload_to='closed_tests/', blank=True, null=True)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        ordering = ['-created_at']
        verbose_name = "Yopiq Test"
        verbose_name_plural = "Yopiq Testlar"

    def __str__(self):
        return self.title

    def update_duration(self):
        """
        FAQAT vaqt ko'rsatkichiga tegmaymiz, agar u 0 bo'lsa yangilaymiz.
        Agar siz qo'lda vaqt kiritsangiz, bu funksiya uni o'zgartirmaydi.
        """
        if self.duration == 0:
            # related_name='closed_questions' ekanligiga ishonch hosil qiling
            count = self.closed_questions.count()
            if count > 0:
                # self.save() qilish o'rniga faqat bitta maydonni yangilash xavfsizroq
                ClosedTest.objects.filter(pk=self.pk).update(duration=count)

    def get_type_label(self):
        method = "Fayldan" if self.file else "Qo'lda"
        return f"Yopiq Test ({method})"


@receiver([post_save, post_delete], sender='amaliyot.ClosedQuestion') # 'amaliyot' - ilovangiz nomi
def auto_update_closed_test_duration(sender, instance, **kwargs):
    """Savol qo'shilsa yoki o'chirilsa, faqat vaqt 0 bo'lsa yangilaydi"""
    if instance.test:
        # Metodni chaqiramiz
        instance.test.update_duration()

  
class ClosedQuestion(models.Model):
    QUESTION_TYPES = [('single', 'Single Choice'), ('multiple', 'Multiple Choice')]

    test = models.ForeignKey('ClosedTest', on_delete=models.CASCADE, related_name='closed_questions')
    text = models.TextField()
    question_type = models.CharField(max_length=20, choices=QUESTION_TYPES, default='single')
    
    
    answer_A = models.TextField()
    answer_B = models.TextField()
    answer_C = models.TextField()
    answer_D = models.TextField()

    is_correct_a = models.BooleanField(default=False)
    is_correct_b = models.BooleanField(default=False)
    is_correct_c = models.BooleanField(default=False)
    is_correct_d = models.BooleanField(default=False)
    
    # To'g'ri javoblar va ixtiyoriy fayl
    correct_answers = models.CharField(max_length=255)  
    file = models.FileField(upload_to='question_files/', blank=True, null=True)

    def __str__(self):
        
        val = self.text if self.text else ""
    
        return str(val)[:50]


class ClosedTestSession(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    test = models.ForeignKey(ClosedTest, on_delete=models.CASCADE)
    started_at = models.DateTimeField(auto_now_add=True)
    completed_at = models.DateTimeField(null=True, blank=True)
    score_obtained = models.FloatField(default=0)

    def __str__(self):
        return f"{self.user.username} - {self.test.title}"


class ClosedAnswer(models.Model):
    session = models.ForeignKey(
        'TestSession', 
        related_name='closed_answers', 
        on_delete=models.CASCADE
    )
    question = models.ForeignKey(ClosedQuestion, on_delete=models.CASCADE)
    selected_answers = models.CharField(max_length=255)
    is_correct = models.BooleanField(default=False)

    def __str__(self):
        user_name = self.session.user.username if self.session.user else "Noma'lum"
        return f"{self.question.text[:30]}"
    

class CaseTest(models.Model):
    # --- 🔗 OTA MODEL BILAN BOG'LANISH ---
    # Bu maydon Search sahifasida 'Kazus' bo'lib ko'rinishi uchun shart!
    parent_test = models.OneToOneField(
        'Test', 
        on_delete=models.CASCADE, 
        related_name='case_test_detail', 
        null=True, 
        blank=True,
        verbose_name="Asosiy test ID"
    )

    # --- 🛠️ 1. ASOSIY MAYDONLAR ---
    title = models.CharField(max_length=255, verbose_name="Test sarlavhasi")
    creator = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Muallif")
    subject = models.ForeignKey('Subject', on_delete=models.CASCADE, verbose_name="Fan nomi")
    
    test_duration = models.PositiveIntegerField(default=0, verbose_name="Test davomiyligi (minut)")
    created_at = models.DateTimeField(auto_now_add=True, null=True, blank=True)

    # --- 🛠️ 2. KAZUSGA XOS MAXSUS MAYDONLAR ---
    case_text = models.TextField(verbose_name="Kazus matni (Muammo bayoni)")
    file = models.FileField(upload_to='case_tests/', null=True, blank=True)
    
    test_type = models.CharField(max_length=20, default='kazus', editable=False)
    creation_method = models.CharField(max_length=20, default='manual')

    class Meta:
        verbose_name = "Kazus testi"
        verbose_name_plural = "Kazus testlar"
        ordering = ['-created_at']

    def save(self, *args, **kwargs):
        # 1. Yaratilish usulini belgilash
        if self.file:
            self.creation_method = 'file'
        
        # 2. Saqlash
        super().save(*args, **kwargs)

        # 3. 🔥 MUHIM: Agar ota model (parent_test) bo'lsa, uning turini ham 'kazus' qilish
        if self.parent_test and self.parent_test.test_type != 'kazus':
            self.parent_test.test_type = 'kazus'
            self.parent_test.save(update_fields=['test_type'])

    def __str__(self):
        return self.title if self.title else f"Kazus #{self.pk}"

    # --- 🛠️ 3. QO'SHIMCHA METODLAR ---
    def get_type_label(self):
        method = "FAYLDAN" if self.creation_method == 'file' else "QO'LDA"
        return f"KAZUS TEST ({method})"
    
    @property
    def toifa(self):
        return self.get_type_label()

    @property
    def questions_count(self):
        if hasattr(self, 'questions'):
            return self.questions.count()
        return 0

    def get_display_name(self):
        """Search sahifasida turni to'g'ri ko'rsatish uchun"""
        return "Kazus Test"


class Olympiad(models.Model):
    name = models.CharField(max_length=255)
    description = models.TextField(blank=True, null=True)
    start_date = models.DateTimeField()
    end_date = models.DateTimeField()
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.name


class OlympiadGroup(models.Model):
    olympiad = models.ForeignKey(
        'Olympiad',
        on_delete=models.CASCADE,
        related_name='groups',
        default=1,
    )
    name = models.CharField(max_length=255, blank=True)  # optional
    test = models.ForeignKey(
        'amaliyot.Test',
        on_delete=models.CASCADE,
        related_name='olympiad_groups',
        null=True,       # optional
        blank=True       # optional
    )
    start_time = models.DateTimeField(null=True, blank=True)  # optional
    end_time = models.DateTimeField(null=True, blank=True)    # optional
    max_participants = models.PositiveIntegerField(null=True, blank=True)
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_olympiads'
    )

    def __str__(self):
        return f"{self.name or 'Olimpiada Group'} ({self.olympiad.name})"
    
    
    def save(self, *args, **kwargs):
        if not self.start_time:
            self.start_time = timezone.now()
        if not self.end_time:
            self.end_time = self.start_time + timezone.timedelta(hours=1)
        super().save(*args, **kwargs)


class OlympiadParticipant(models.Model):
    group = models.ForeignKey('OlympiadGroup', on_delete=models.CASCADE, related_name="participants")
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name="olympiad_participations")
    joined_at = models.DateTimeField(auto_now_add=True)
    start_time = models.DateTimeField(blank=True, null=True)  
    score = models.FloatField(default=0)
    finished = models.BooleanField(default=False)
    correct_answers = models.PositiveIntegerField(default=0)
    wrong_answers = models.PositiveIntegerField(default=0)
    last_activity = models.DateTimeField(auto_now=True)

    class Meta:
        unique_together = ('group', 'user')

    def __str__(self):
        
        try:
            return f"{self.user.userprofile.full_name} - {self.group.name}"
        except:
            return f"{self.user.username} - {self.group.name}"


class OlympiadAnswer(models.Model):
    participant = models.ForeignKey(OlympiadParticipant, on_delete=models.CASCADE, related_name="answers")
    question_id = models.PositiveIntegerField()  
    selected = models.CharField(max_length=255, blank=True, null=True)
    is_correct = models.BooleanField(default=False)
    answered_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        unique_together = ("participant", "question_id")

    def __str__(self):
        return f"Answer by {self.participant.user.username} (Q{self.question_id})"
    

class OlympiadQuestion(models.Model):
    group = models.ForeignKey(OlympiadGroup, on_delete=models.CASCADE, related_name="questions")
    text = models.TextField()
    options = models.JSONField(null=True, blank=True)   
    difficulty = models.PositiveIntegerField(default=1)
    created_at = models.DateTimeField(auto_now_add=True)
    correct_answer = models.CharField(max_length=10, null=True, blank=True) 


class Group(models.Model):
    name = models.CharField(max_length=100, verbose_name="Guruh nomi")
    creator = models.ForeignKey(User, on_delete=models.CASCADE, related_name='created_groups')
    created_at = models.DateTimeField(default=timezone.now)

    start_time = models.DateTimeField(null=True, blank=True)
    end_time = models.DateTimeField(null=True, blank=True)
    max_participants = models.PositiveIntegerField(default=0)
    
    invite_token = models.UUIDField(default=uuid.uuid4, editable=False, unique=True)
    
    def __str__(self):
        return self.name
      

class GroupMembership(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='group_memberships')
    group = models.ForeignKey(Group, on_delete=models.CASCADE, related_name='memberships')
    joined_at = models.DateTimeField(default=timezone.now)

    class Meta:
        unique_together = ('user', 'group')  

    def __str__(self):
        return f"{self.user.username} → {self.group.name}"
    

class ChatMessage(models.Model):
    group = models.ForeignKey(Group, on_delete=models.CASCADE, related_name='messages')
    sender = models.ForeignKey(User, on_delete=models.CASCADE, related_name='sent_messages')
    content = models.TextField()
    timestamp = models.DateTimeField(auto_now_add=True)

    class Meta:
        
        ordering = ['timestamp']

    def __str__(self):
        return f"[{self.group.name}] {self.sender.username}: {self.content[:50]}..."
    

def generate_random_key():
    return uuid.uuid4().hex[:10].upper()


class Subject(models.Model):
    name = models.CharField(max_length=255)

    def __str__(self):
        return self.name


class Topic(models.Model):
    name = models.CharField(max_length=100)
    subject = models.ForeignKey(Subject, on_delete=models.CASCADE, related_name='topics')

    def __str__(self):
        return self.name
    

class Choice(models.Model):
    text = models.CharField(max_length=255)  
    is_correct = models.BooleanField(default=False)  

    def __str__(self):
        return self.text


class Answer(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='answers')  
    question = models.ForeignKey('Question', on_delete=models.CASCADE, related_name='answers') 
    selected_choices = models.ManyToManyField(Choice, related_name='answers')  
    answered_at = models.DateTimeField(auto_now_add=True)  

    def __str__(self):
        return f'Answer by {self.user} to Question {self.question.id}'
    

class UploadedFile(models.Model):
    file = models.FileField(upload_to='uploaded_tests/')
    test = models.ForeignKey('Test', on_delete=models.CASCADE, related_name='uploaded_files', null=True, blank=True)
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"{self.file.name} ({self.test.title if self.test else 'Test yaratilmagan'})"


class Test(models.Model):
    VIEW_MODE_CHOICES = [
        ('sequential', 'Sequential'),
        ('random', 'Random'),
    ]
    TEST_TYPE_CHOICES = [
        ('ordinary', 'Oddiy Test'),
        ('kazus', 'Kazus Test'),
    ]
    CREATION_METHOD_CHOICES = [
        ('manual', 'Qo\'lda yaratilgan'),
        ('file', 'Fayldan yuklangan'),
    ]

    title = models.CharField(max_length=255)
    category = models.CharField(max_length=100)
    test_type = models.CharField(max_length=20, choices=TEST_TYPE_CHOICES, default='ordinary')
    creation_method = models.CharField(max_length=20, choices=CREATION_METHOD_CHOICES, default='manual')
    is_public = models.BooleanField(default=True)
    created_at = models.DateTimeField(auto_now_add=True)
    description = models.TextField(blank=True)
    subject = models.ForeignKey('Subject', on_delete=models.CASCADE)
    creator = models.ForeignKey(User, on_delete=models.CASCADE)
    is_paid = models.BooleanField(default=False)
    price = models.IntegerField(null=True, blank=True)
    access_key = models.CharField(max_length=20, unique=True, null=True, blank=True)
    view_mode = models.CharField(max_length=50, choices=VIEW_MODE_CHOICES, default='sequential')
    
    # Default 0 bo'lib qoladi, bu bizga "vaqt hali belgilanmagan" degan signalni beradi
    test_duration= models.PositiveIntegerField(default=0, verbose_name="Vaqt (minut)")
    score = models.IntegerField(default=100) 
    is_multiple_choice = models.BooleanField(default=False)


    def get_display_name(self):
        method = "Fayldan" if self.creation_method == 'file' else "Qo'lda"
        
        if self.test_type == 'kazus':
            return f"Kazus Test ({method})"
        if self.test_type == 'closed':
            return f"Yopiq Test ({method})"
            
        return f"Oddiy Test ({method})"

     # Admin panelda sarlavha o'rnida ham chiqishi uchun
    def __str__(self):
        return f"{self.title} | {self.get_display_name()}"


    def get_real_label(self):
       # Eng muhim qism: related_name orqali tekshiramiz
       if hasattr(self, 'case_test_detail'):
           return "KAZUS TEST"
       if hasattr(self, 'closedtest') or hasattr(self, 'closed_test'):
           return "YOPIQ TEST"
       return "ODDIY TEST"

    def get_full_label(self):
        is_file = False
        # CaseTest yoki oddiy testligidan qat'iy nazar faylni tekshiramiz
        if hasattr(self, 'case_test_detail') and self.case_test_detail.file:
            is_file = True
        elif self.creation_method == 'file':
            is_file = True
        
        label = self.get_real_label() # KAZUS TEST, YOPIQ TEST yoki ODDIY TEST
        method = "(FAYLDAN)" if is_file else "(QO'LDA)"
        return f"{label} {method}"

    def save(self, *args, **kwargs):
        # 1. Pulli testlar mantiqi
        if self.is_paid:
            if self.price is None:
                raise ValueError("Pulli testlar uchun narx kiritish shart.")
            if not self.access_key:
                self.access_key = get_random_string(10).upper()
        else:
            self.price = None
            self.access_key = None
        
        if self.pk:
            try:
                count = self.questions.count()
                if count > 0 and self.test_duration == 0:
                    self.test_duration = count
            except Exception:
                pass

        super().save(*args, **kwargs)

    def update_duration_minutes(self):
        if self.pk:
            count = self.questions.count()
            if count > 0:
                # Test.objects.filter(id=self.id) o'rniga:
                Test.objects.filter(pk=self.pk).update(test_duration=count) # test_duration maydon nomiga e'tibor bering
                self.test_duration = count

    def __str__(self):
        return self.title
    
    def get_type_label(self):
        # 1. Avval 'test_type' maydonini tekshiramiz
        # 2. Keyin CaseTest bog'liqligi borligini tekshiramiz (case_test_detail - sizning modelda)
        if self.test_type == 'kazus' or hasattr(self, 'case_test_detail'):
            t_name = "Kasuz Test"
        elif self.test_type == 'closed' or hasattr(self, 'closed_test'):
            t_name = "Yopiq Test"
        else:
            t_name = "Oddiy Test"

        method = "Fayldan" if self.creation_method == 'file' else "Qo'lda"
        return f"{t_name}({method})"
    

class Result(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='results')
    test = models.ForeignKey(Test, on_delete=models.CASCADE, null=True, blank=True)
    total = models.IntegerField()
    percentage = models.FloatField()
    date = models.DateTimeField(auto_now_add=True)
    group = models.ForeignKey(Group, on_delete=models.SET_NULL, null=True, blank=True)
    score = models.FloatField(default=0)
    correct_count = models.IntegerField(default=0)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        ordering = ['-score']

    def __str__(self):
        return f"{self.user.username} - {self.score}"
    

class CaseQuestion(models.Model):
    # --- 🔗 BOG'LIQLIK ---
    case_test = models.ForeignKey('CaseTest', on_delete=models.CASCADE, related_name='questions')
    
    # --- 📝 SAVOL VA VARIANTLAR ---
    text = models.TextField(verbose_name="Savol matni")
    # Kazus matni uchun maydon (agar kodingizda ishlatilgan bo'lsa)
    case_text = models.TextField(blank=True, default="", verbose_name="Kazus sharti")
    
    
    answer_a = models.CharField(max_length=255)
    answer_b = models.CharField(max_length=255)
    answer_c = models.CharField(max_length=255, blank=True, null=True)
    answer_d = models.CharField(max_length=255, blank=True, null=True)
    
    # To'g'ri javoblar uchun yordamchi maydonlar
    is_correct_a = models.BooleanField(default=False)
    is_correct_b = models.BooleanField(default=False)
    is_correct_c = models.BooleanField(default=False)
    is_correct_d = models.BooleanField(default=False)
    
    # JSON formatda to'g'ri javoblar (masalan: ["A", "C"])
    correct_answers = models.JSONField(default=list)
    
    # Qo'shimcha flaglar
    is_multiple_choice = models.BooleanField(default=False)
    question_type = models.CharField(max_length=20, default='case')

    # --- 📁 FAYL BILAN ISHLASH ---
    file = models.FileField(upload_to='case_question_files/', null=True, blank=True)

    class Meta:
        verbose_name = "Kazus savoli"
        verbose_name_plural = "Kazus savollari"

    def __str__(self):
        return f"Savol: {self.text[:50]}..."

    def save(self, *args, **kwargs):
        # 1. Word Fayldan ma'lumotlarni o'qish
        if self.file:
            try:
                self.file.seek(0)
                file_content = self.file.read()
                
                doc = Document(io.BytesIO(file_content))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                full_text = "\n".join(paragraphs)
                
                if full_text:
                    # TO'SIQ: Agar faylda [k] bo'lmasa, Kazus modeliga saqlamaymiz
                    if '[k]' not in full_text:
                        print("❌ Bu oddiy savol! CaseQuestion modeliga saqlanmadi.")
                        return 

                    # Faylni [k] bo'yicha bo'laklarga bo'lish
                    question_blocks = re.split(r'\[k\]', full_text)
                    blocks = [b.strip() for b in question_blocks if b.strip()]
                    
                    if blocks:
                        current_block = blocks[0] 
                        lines = [l.strip() for l in current_block.split('\n') if l.strip()]
                        
                        # Parser mantiqi: lines[0]=Kazus, lines[1]=Savol, lines[2:]=Variantlar
                        if lines and len(lines) > 2:
                            self.case_text = lines[0] # Kazus sharti
                            self.text = lines[1]      # Savol matni
                            
                            options = []
                            temp_correct_answers = []
                            char_map = ['A', 'B', 'C', 'D']
                            
                            for line in lines[2:]:
                                clean_line = line
                                is_correct = False
                                
                                if line.startswith('+'):
                                    is_correct = True
                                    clean_line = line[1:].strip()
                                elif line.startswith('-'):
                                    clean_line = line[1:].strip()

                                if len(options) < 4:
                                    lbl = char_map[len(options)]
                                    options.append(clean_line)
                                    if is_correct:
                                        temp_correct_answers.append(lbl)

                            # Variantlarni modelga yuklash
                            self.answer_a = options[0] if len(options) > 0 else ""
                            self.answer_b = options[1] if len(options) > 1 else ""
                            self.answer_c = options[2] if len(options) > 2 else ""
                            self.answer_d = options[3] if len(options) > 3 else ""
                            self.correct_answers = temp_correct_answers

            except Exception as e:
                print(f"❌ CaseQuestion Word Parser Xatosi: {str(e)}")

        # 2. MA'LUMOTLARNI QAT'IY SINXRONLASH
        if not isinstance(self.correct_answers, list):
            self.correct_answers = []

        # Admin paneldagi BooleanField (is_correct_x) lar bilan sinxronlash
        current_flags = []
        if self.is_correct_a: current_flags.append('A')
        if self.is_correct_b: current_flags.append('B')
        if self.is_correct_c: current_flags.append('C')
        if self.is_correct_d: current_flags.append('D')

        if current_flags:
            # Agar checkboxlar belgilangan bo'lsa, JSON ni yangilaymiz
            self.correct_answers = current_flags 
        
        # Checkboxlarni JSON dagi ma'lumotga qarab to'g'irlaymiz
        self.is_correct_a = ('A' in self.correct_answers)
        self.is_correct_b = ('B' in self.correct_answers)
        self.is_correct_c = ('C' in self.correct_answers)
        self.is_correct_d = ('D' in self.correct_answers)

        # 3. SAVOL TURI
        self.question_type = 'case'
        self.is_multiple_choice = len(self.correct_answers) > 1

        super().save(*args, **kwargs)
    

class Question(models.Model):
    QUESTION_TYPE = [
        ('single', 'Single Choice'),
        ('multiple', 'Multiple Choice'),
        ('case', 'Case Study (Kazus)'),
    ]

    # 🔹 Bog'liqliklar
    test = models.ForeignKey('Test', on_delete=models.CASCADE, related_name='questions', null=True, blank=True)
    closed_test = models.ForeignKey('ClosedTest', on_delete=models.CASCADE, null=True, blank=True, related_name='questions')
    case_test = models.ForeignKey('CaseTest', on_delete=models.CASCADE, null=True, blank=True, related_name='case_questions')
    subject = models.ForeignKey('Subject', on_delete=models.SET_NULL, null=True, blank=True)
    topic = models.ForeignKey('Topic', on_delete=models.SET_NULL, null=True, blank=True)
    creator = models.ForeignKey(settings.AUTH_USER_MODEL, on_delete=models.CASCADE, related_name='questions')
    group = models.ForeignKey('Group', on_delete=models.SET_NULL, null=True, blank=True)

    # 🔹 Asosiy maydonlar
    difficulty = models.CharField(max_length=50, blank=True, default="")
    text = models.TextField()
    case_text = models.TextField(blank=True, default="", help_text="Kazusning umumiy sharti")
    image = models.ImageField(upload_to='question_images/', null=True, blank=True)
    youtube_link = models.URLField(blank=True, default="")
    file = models.FileField(upload_to='question_files/', null=True, blank=True)

    # 🔹 To'lov va Kirish kaliti
    is_paid = models.BooleanField(default=False)
    price = models.PositiveIntegerField(null=True, blank=True)
    access_key = models.CharField(max_length=20, unique=True, null=True, blank=True)

    # 🔹 Javoblar mantig'i
    correct_answers = models.JSONField(default=list, help_text="To'g'ri javoblar ro'yxati (masalan: ['A', 'C'])")
    correct_answer = models.CharField(max_length=1, null=True, blank=True, help_text="Asosiy to'g'ri javob (A, B, C yoki D)")
    correct_option = models.CharField(max_length=1, blank=True) 

    # 🔹 Savol turi
    question_type = models.CharField(max_length=10, choices=QUESTION_TYPE, default='single')
    is_multiple_choice = models.BooleanField(default=False)

    # 🔹 Variantlar (A, B, C, D)
    answer_a = models.CharField(max_length=255, blank=True, default="")
    answer_b = models.CharField(max_length=255, blank=True, default="")
    answer_c = models.CharField(max_length=255, blank=True, default="")
    answer_d = models.CharField(max_length=255, blank=True, default="")

    is_correct_a = models.BooleanField(default=False)
    is_correct_b = models.BooleanField(default=False)
    is_correct_c = models.BooleanField(default=False)
    is_correct_d = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.text[:50]}..." if self.text else "Savol matni yo'q"

    def save(self, *args, **kwargs):
        # 1. Word Fayldan ma'lumotlarni o'qish
        if self.file:
            try:
                self.file.seek(0)
                file_content = self.file.read()
                
                doc = Document(io.BytesIO(file_content))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                full_text = "\n".join(paragraphs)
                
                if full_text:
                    # TO'SIQ: Agar bu kazus fayli bo'lsa, oddiy modelga saqlamaymiz
                    if '[k]' in full_text:
                        print("❌ Bu Kazus fayli! Oddiy Question modeliga saqlanmadi.")
                        return 

                    # Faqat [s] bo'yicha split qilamiz
                    question_blocks = re.split(r'\[s\]', full_text)
                    blocks = [b.strip() for b in question_blocks if b.strip()]
                    
                    if blocks:
                        current_block = blocks[0] 
                        lines = [l.strip() for l in current_block.split('\n') if l.strip()]
                        
                        if lines:
                            self.text = lines[0] # Savol matni
                            start_idx = 1
                            
                            options = []
                            temp_correct_answers = []
                            char_map = ['A', 'B', 'C', 'D']
                            
                            for line in lines[start_idx:]:
                                clean_line = line
                                is_correct = False
                                if line.startswith('+'):
                                    is_correct = True
                                    clean_line = line[1:].strip()
                                elif line.startswith('-'):
                                    clean_line = line[1:].strip()

                                if len(options) < 4:
                                    lbl = char_map[len(options)]
                                    options.append(clean_line)
                                    if is_correct:
                                        temp_correct_answers.append(lbl)

                            self.answer_a = options[0] if len(options) > 0 else ""
                            self.answer_b = options[1] if len(options) > 1 else ""
                            self.answer_c = options[2] if len(options) > 2 else ""
                            self.answer_d = options[3] if len(options) > 3 else ""
                            self.correct_answers = temp_correct_answers
            except Exception as e:
                print(f"❌ Word Parser Error: {str(e)}")

        # 2. MA'LUMOTLARNI QAT'IY SINXRONLASH
        if not isinstance(self.correct_answers, list):
            self.correct_answers = []

        # JORIIY HOLATNI ANIQLASH
        current_flags = []
        if self.is_correct_a: current_flags.append('A')
        if self.is_correct_b: current_flags.append('B')
        if self.is_correct_c: current_flags.append('C')
        if self.is_correct_d: current_flags.append('D')

        final_letter = ""

        # A) Checkboxlar ustuvor
        if current_flags:
            final_letter = current_flags[0]
            self.correct_answers = current_flags 
        
        # B) correct_answer (harf)
        elif self.correct_answer and self.correct_answer.strip():
            final_letter = self.correct_answer.upper().strip()
            self.correct_answers = [final_letter]

        # C) correct_answers (list)
        elif self.correct_answers:
            final_letter = str(self.correct_answers[0]).upper()

        # D) correct_option
        elif self.correct_option and self.correct_option.strip():
            final_letter = self.correct_option.upper().strip()
            self.correct_answers = [final_letter]

        # Barcha maydonlarni bir xillashtirish
        if final_letter in ['A', 'B', 'C', 'D']:
            self.correct_answer = final_letter
            self.correct_option = final_letter
            self.is_correct_a = ('A' in self.correct_answers)
            self.is_correct_b = ('B' in self.correct_answers)
            self.is_correct_c = ('C' in self.correct_answers)
            self.is_correct_d = ('D' in self.correct_answers)
        else:
            self.correct_answer = None
            self.correct_option = ""

        # 3. SAVOL TURI VA SAQLASH
        self.is_multiple_choice = len(self.correct_answers) > 1
        if self.question_type != 'case':
            self.question_type = 'multiple' if self.is_multiple_choice else 'single'

        super().save(*args, **kwargs)


class TestQueue(models.Model):
    STATUS_CHOICES = (
        ('waiting', 'Waiting'),
        ('active', 'Active'),
        ('done', 'Done')
    )

    user = models.ForeignKey(User, on_delete=models.CASCADE)
    test = models.ForeignKey('Test', on_delete=models.CASCADE)
    start_time = models.DateTimeField()
    expected_end_time = models.DateTimeField()
    score = models.FloatField(null=True, blank=True)
    stopped_by_user = models.BooleanField(default=False)
    is_completed = models.BooleanField(default=False)
    current_step = models.IntegerField(default=0)
    status = models.CharField(max_length=10, choices=STATUS_CHOICES, default='waiting')

    def __str__(self):
        return f"User: {self.user.username} - Test: {self.test.title} - Step: {self.current_step}"

    def is_time_expired(self):
        return timezone.now() >= self.expected_end_time and not self.is_completed

    def auto_finish(self):
        """Tizim avtomatik yakunlaydi agar foydalanuvchi o‘zi yakunlamagan bo‘lsa"""
        if self.is_time_expired():
            self.status = 'done'
            self.is_completed = True
            self.stopped_by_user = False
            self.save()


class QuestionChoice(models.Model):
    question = models.ForeignKey(Question, on_delete=models.CASCADE, related_name='question_choices')
    text = models.CharField(max_length=255)
    is_correct = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.text} ({'Correct' if self.is_correct else 'Incorrect'})"


class Option(models.Model):
    question = models.ForeignKey(Question, on_delete=models.CASCADE, related_name='options')
    text = models.CharField(max_length=255)
    is_correct = models.BooleanField(default=False)

    def __str__(self):
        return self.text


class QuestionImage(models.Model):
    question = models.ForeignKey(Question, on_delete=models.CASCADE, related_name='images')
    image = models.ImageField(upload_to='question_images/')

    def __str__(self):
        return f"Rasm {self.id} - {self.image.name}"
    

class TestSession(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    test = models.ForeignKey('Test', on_delete=models.CASCADE, null=True, blank=True)
    closed_test = models.ForeignKey('ClosedTest', on_delete=models.CASCADE, null=True, blank=True)
    # --- YANGI QATOR ---
    case_test = models.ForeignKey('CaseTest', on_delete=models.CASCADE, null=True, blank=True)
    
    queue = models.ForeignKey('TestQueue', on_delete=models.CASCADE, null=True)
    current_index = models.IntegerField(default=0)
    auto_evaluated = models.BooleanField(default=False) 
    started_at = models.DateTimeField(default=timezone.now)
    finished_at = models.DateTimeField(null=True, blank=True)
    start_time = models.DateTimeField(default=timezone.now)
    end_time = models.DateTimeField(blank=True, null=True)
    is_finished = models.BooleanField(default=False)
    session_id = models.UUIDField(default=uuid.uuid4, unique=True, editable=False)
    score = models.IntegerField(null=True, blank=True)
    total_questions = models.IntegerField(default=0)
    correct_answers = models.IntegerField(default=0)
    incorrect_answers = models.IntegerField(default=0)
    payment_amount = models.DecimalField(max_digits=10, decimal_places=2, null=True, blank=True)

    def __str__(self):
        if self.test:
            title = self.test.title
        elif self.closed_test:
            title = self.closed_test.title
        elif self.case_test:
            title = self.case_test.title
        else:
            title = "Noma'lum"
        return f"{self.user.username} - {title}"

    def get_calculated_duration(self):
        """Har bir savol uchun vaqt hisoblash mantiqi"""
        # 1. Kazus test bo'lsa
        if self.case_test:
            # CaseTest modelida 'test_duration' bor bo'lsa shuni oladi, bo'lmasa savollar soni
            duration = getattr(self.case_test, 'test_duration', 0)
            if not duration:
                duration = self.case_test.questions.count()
            return duration if duration > 0 else 15

        # 2. Oddiy test bo'lsa
        if self.test:
            q_count = self.test.questions.count()
            return q_count if q_count > 0 else 1

        # 3. Yopiq test bo'lsa
        if self.closed_test:
            q_count = self.closed_test.closed_questions.count()
            return q_count if q_count > 0 else 1

        return 1

    # --- TAYMER METODLARI O'ZGARMAYDI ---
    def has_time_expired(self):
        duration = self.get_calculated_duration()
        expire_time = self.start_time + timezone.timedelta(minutes=duration)
        return timezone.now() > expire_time

    def time_left_seconds(self):
        elapsed = (timezone.now() - self.start_time).total_seconds()
        duration = self.get_calculated_duration()
        return max(0, (duration * 60) - elapsed)

    def save(self, *args, **kwargs):
        if not self.pk:
            if self.test:
                self.total_questions = self.test.questions.count()
            elif self.closed_test:
                if hasattr(self.closed_test, 'closed_questions'):
                    self.total_questions = self.closed_test.closed_questions.count()
                else:
                    self.total_questions = 0
            # --- YANGI KAZUS MANTIQI ---
            elif self.case_test:
                self.total_questions = self.case_test.questions.count()
        
        super().save(*args, **kwargs)
    

class TestAnswer(models.Model):
    session = models.ForeignKey(TestSession, related_name='test_answers', on_delete=models.CASCADE)
    user = models.ForeignKey(User, on_delete=models.CASCADE, null=True, blank=True)
    question = models.ForeignKey(Question, on_delete=models.CASCADE)
    selected_answers = models.JSONField(default=list)
    created_at = models.DateTimeField(auto_now_add=True)

    class Meta:
        unique_together = ('session', 'question')  

    def __str__(self):
        return f"Session {self.session.id} - Question {self.question.id}"
    

class UserAnswer(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    session = models.ForeignKey(TestSession, on_delete=models.CASCADE, related_name='user_answers')
    question = models.ForeignKey(Question, on_delete=models.CASCADE)
    selected_options = models.ManyToManyField(Option)
    is_correct = models.BooleanField(default=False)
    answered_at = models.DateTimeField(default=timezone.now)

    def __str__(self):
        return f"{self.user.username} - {self.question.text[:30]} - {'to‘g‘ri' if self.is_correct else 'noto‘g‘ri'}"


class Payment(models.Model):
    PAYMENT_PROVIDER_CHOICES = [
        ('click', 'Click'),
        ('payme', 'Payme'),
    ]

    PAYMENT_STATUS_CHOICES = [
        ('pending', 'Kutilmoqda'),
        ('success', 'Muvaffaqiyatli'),
        ('failed', 'Muvaffaqiyatsiz'),
    ]

    user = models.ForeignKey(
        User, on_delete=models.CASCADE, related_name='payments'
    )
    test = models.ForeignKey(
        'Test', on_delete=models.CASCADE, null=True, blank=True, related_name='payments'
    )
    question = models.ForeignKey(
        'Question', on_delete=models.CASCADE, null=True, blank=True
    )

    provider = models.CharField(
        max_length=10, choices=PAYMENT_PROVIDER_CHOICES
    )
    transaction_id = models.CharField(
        max_length=255, blank=True, null=True, db_index=True  # ✅ null=True qo‘shildi
    )
    amount = models.DecimalField(
        max_digits=10, decimal_places=2, default=Decimal('0.00')
    )

    status = models.CharField(
        max_length=10, choices=PAYMENT_STATUS_CHOICES, default='pending'
    )
    is_successful = models.BooleanField(default=False)
    reason = models.TextField(blank=True, null=True)

    created_at = models.DateTimeField(auto_now_add=True)
    pay_time = models.DateTimeField(null=True, blank=True)

    def mark_successful(self, transaction_id=None):
        """To‘lov muvaffaqiyatli tugaganini belgilaydi."""
        self.status = 'success'
        self.is_successful = True
        if transaction_id:
            self.transaction_id = transaction_id
        self.pay_time = timezone.now()
        self.save(update_fields=['status', 'is_successful', 'transaction_id', 'pay_time'])

    def mark_failed(self, reason=None):
        """To‘lov muvaffaqiyatsiz bo‘lganda belgilaydi."""
        self.status = 'failed'
        self.is_successful = False
        self.reason = reason
        self.save(update_fields=['status', 'is_successful', 'reason'])

    def __str__(self):
        return f"{self.user.username} | {self.get_provider_display()} | {self.amount} so‘m | {self.get_status_display()}"

    class Meta:
        ordering = ['-created_at']
        verbose_name = "To‘lov"
        verbose_name_plural = "To‘lovlar"


class TestCompletion(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    test = models.ForeignKey(Test, on_delete=models.CASCADE)
    score = models.IntegerField()
    is_completed = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.test.title} - {self.user.username}"


class UserTest(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    test = models.ForeignKey(Test, on_delete=models.CASCADE)
    completed = models.BooleanField(default=False)
    score = models.IntegerField(null=True, blank=True)

    def __str__(self):
        return f"{self.user.username} - {self.test.title}"


def user_avatar_upload_path(instance, filename):
    return f'avatars/user_{instance.user.id}/{filename}'


class UserProfile(models.Model):
    GENDER_CHOICES = (
        ('male', "O‘g‘il bola"),
        ('female', "Qiz bola"),
    )

    user = models.OneToOneField(User, on_delete=models.CASCADE, related_name='userprofile')
    first_name = models.CharField(max_length=30, blank=True)
    last_name = models.CharField(max_length=30, blank=True)
    birth_date = models.DateField(null=True, blank=True)
    bio = models.TextField(blank=True)
    location = models.CharField(max_length=255, blank=True)
    gender = models.CharField(max_length=10, choices=GENDER_CHOICES, blank=True, null=True)
    avatar = models.ImageField(
        upload_to="avatars/", # user_avatar_upload_path funksiyangiz bo'lsa o'shani yozing
        blank=True,
        null=True
    )
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def __str__(self):
        return f"{self.full_name} ({self.user.username})"

    @property
    def full_name(self):
        # Ism va familiyani birlashtirish
        name = f"{self.first_name} {self.last_name}".strip()
        if name:
            return name
        
        # Agar profile bo'sh bo'lsa, Django User modelidagi ismni tekshirish
        django_name = self.user.get_full_name().strip()
        if django_name:
            return django_name
            
        # Hammasi bo'sh bo'lsa, login (username) qaytarish
        return self.user.username

    @property
    def avatar_url(self):
        if self.avatar:
            return self.avatar.url
        default_file = "defaults/female.png" if self.gender == "female" else "defaults/male.png"
        return settings.MEDIA_URL + default_file
    

@receiver(post_save, sender=User)
def create_or_update_user_profile(sender, instance, created, **kwargs):
    if created:
        
        UserProfile.objects.get_or_create(
            user=instance,
            defaults={
                'first_name': instance.first_name,
                'last_name': instance.last_name
            }
        )


class Task(models.Model):
    name = models.CharField(max_length=100)
    status = models.CharField(max_length=50)
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    completed = models.BooleanField(default=False)


class ApiKey(models.Model):
    user = models.ForeignKey(User, on_delete=models.CASCADE)
    key = models.CharField(max_length=100, unique=True)
    created_at = models.DateTimeField(auto_now_add=True)
    is_active = models.BooleanField(default=True)
    expires_at = models.DateTimeField(null=True, blank=True)

    def __str__(self):
        return f"{self.user.username} - {self.key}"









